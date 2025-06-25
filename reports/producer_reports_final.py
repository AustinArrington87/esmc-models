import pandas as pd
from docx import Document
from docx.shared import Inches, Pt
from docx.oxml import OxmlElement
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.text import WD_UNDERLINE
from docx.shared import RGBColor
from docx.oxml.ns import qn
import os

# year generating reports for
run_year = "2023"
# used for mapping anonymized producer name 
mapping_df = pd.read_csv('Mapping.csv')
producer_mapping = dict(zip(mapping_df['project_producer'], mapping_df['display_name']))

def load_data():
    # Load carbon data
    carbon_data = pd.read_csv('2023_Verified.csv')
    
    # Load water impacts data
    water_data = pd.read_excel('2024_PLET_Outcomes.xlsx')
    
    return carbon_data, water_data

def format_cell_data(data):
    try:
        float_data = float(data)
        return f"{float_data:.3f}"
    except ValueError:
        return data

def create_table(doc, data, headers, font_name, font_size, font_color, scale_emissions=False):
    table = doc.add_table(rows=len(data) + 1, cols=len(headers))
    table.style = 'Table Grid'

    for i, header in enumerate(headers):
        cell = table.cell(0, i)
        run = cell.paragraphs[0].add_run(header)
        run.font.name = font_name
        run.font.size = font_size
        run.font.color.rgb = font_color
        run.font.bold = True
        cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    for row_idx, row_data in enumerate(data, start=1):
        for col_idx, cell_data in enumerate(row_data):
            if scale_emissions and isinstance(cell_data, (int, float)):
                cell_data *= 0.001
            if isinstance(cell_data, float):
                cell_data = f"{cell_data:.3f}"
            cell = table.cell(row_idx, col_idx)
            run = cell.paragraphs[0].add_run(str(cell_data))
            run.font.name = font_name
            run.font.size = font_size
            run.font.color.rgb = font_color
            cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

def calculate_delta(group, baseline_column, practice_column):
    return group[baseline_column] - group[practice_column]

def add_custom_bullet(doc, label, value):
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(label)
    run.bold = True
    run.font.name = 'Calibri'
    run.font.size = Pt(12)

    value_run = p.add_run(value)
    value_run.font.name = 'Calibri'
    value_run.font.size = Pt(12)

def add_hyperlink(paragraph, text, url):
    part = paragraph.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)

    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)
    hyperlink.set(qn('w:history'), '1')

    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    c = OxmlElement('w:color')
    c.set(qn('w:val'), '0000FF')
    rPr.append(c)
    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'single')
    rPr.append(u)
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)

    return hyperlink

def initialize_document():
    doc = Document()
    # No header image per new requirements
    return doc

def generate_report(carbon_data, water_data, doc):
    font_name = 'Calibri'
    font_size = Pt(12)
    font_color = RGBColor(0, 0, 0)
    blue_color = RGBColor(0, 112, 192)
    green_color = RGBColor(112, 173, 71)

    for producer, group in carbon_data.groupby('Producer (Project)'):
        if producer not in doc.producers:
            doc.producers[producer] = initialize_document()

        document = doc.producers[producer]
        document.add_paragraph()
        # Title in blue
        para = document.add_paragraph()
        run = para.add_run(f"2025 Eco-Harvest Report – ")
        run.font.name = font_name
        run.font.size = Pt(24)
        run.font.bold = True
        run.font.underline = True
        run.font.color.rgb = blue_color
        # Producer Name in blue
        display_name = producer_mapping.get(producer, producer)
        run2 = para.add_run(f"({display_name})")
        run2.font.name = font_name
        run2.font.size = Pt(24)
        run2.font.bold = True
        run2.font.underline = True
        run2.font.color.rgb = blue_color

        # Earnings/intro paragraph
        intro_para = document.add_paragraph()
        run_intro = intro_para.add_run("You earned $XX in the Eco-Harvest program. Thank you for participating and for your dedication to regenerative agriculture.")
        run_intro.font.name = font_name
        run_intro.font.size = font_size
        run_intro.font.color.rgb = font_color

        # Your Results section (green header)
        para = document.add_paragraph()
        run = para.add_run("Your Results")
        run.font.name = font_name
        run.font.size = Pt(16)
        run.font.bold = True
        run.font.underline = True
        run.font.color.rgb = green_color

        # Calculate carbon impacts
        total_reductions = group['reduced'].sum()
        total_removals = group['removed'].sum()
        total_carbon_g = total_reductions + total_removals
        total_acres = group['Acres'].sum()

        add_custom_bullet(document, "Total Carbon Impacts Generated: ", f"{total_carbon_g:.3f} metric tonnes of carbon dioxide equivalent (mtCO2e)")
        add_custom_bullet(document, "Emission Reductions: ", f"Your practice changes avoided {total_reductions:.3f} mtCO2e.")
        add_custom_bullet(document, "Soil Carbon Removals: ", f"Your soils stored {total_removals:.3f} mtCO2e")
        add_custom_bullet(document, "Total Modeled Acres: ", f"{total_acres:.3f}")

        # Add asterisk note for modeled acres
        note_para = document.add_paragraph()
        note_run = note_para.add_run("*ESMC requires complete historical data to generate a baseline, so fields with historical data gaps are not modeled.")
        note_run.font.name = font_name
        note_run.font.size = Pt(9)

        # Carbon Impact Explained section
        para = document.add_paragraph()
        run = para.add_run("Carbon Impact Explained")
        run.font.name = font_name
        run.font.size = Pt(14)
        run.font.bold = True
        run.font.underline = True
        run.font.color.rgb = green_color
        document.add_paragraph("ESMC uses historical data, a scientific model, soil samples and weather data to calculate emissions reductions and soil carbon.")
        add_custom_bullet(document, "Greenhouse Gas Emissions (GHG) Reductions: ", "are the sum of the reductions of methane (CH4), the gas form of nitrogen (N2O) and supply chain and operational emissions**.")
        add_custom_bullet(document, "Soil Carbon: ", "The difference in the amount of carbon in the soil.")

        # Add double asterisk note for supply chain and operational emissions
        note2_para = document.add_paragraph()
        note2_run = note2_para.add_run("**supply chain and operational emissions are the emissions from upstream and on-field activities such as fuel used on farm and imbedded emissions of fertilizer products; ESMC utilized emission factors from Ecoinvent and WFLD")
        note2_run.font.name = font_name
        note2_run.font.size = Pt(9)

        # Impact Breakdown Table
        # Calculate deltas for each field and sum
        direct_n = (group['baseline_n20_direct'] - group['n2o_direct']).sum() if 'baseline_n20_direct' in group and 'n2o_direct' in group else 0
        indirect_n = (group['baseline_n2o_indirect'] - group['n2o_indirect']).sum() if 'baseline_n2o_indirect' in group and 'n2o_indirect' in group else 0
        methane = (group['baseline_methane'] - group['methane']).sum() if 'baseline_methane' in group and 'methane' in group else 0
        # Convert supply chain emissions from kgCO2e to metric tonnes CO2e
        supply_chain = ((group['field_baseline_emissions'] - group['field_practice_emissions']).sum() / 1000) if 'field_baseline_emissions' in group and 'field_practice_emissions' in group else 0
        total_ghg = direct_n + indirect_n + methane + supply_chain
        soil_carbon = (group['dsoc'] - group['baseline_dsoc']).sum() if 'dsoc' in group and 'baseline_dsoc' in group else 0
        total_impact = total_ghg + soil_carbon
        impact_data = [
            ["Direct Nitrogen", f"{direct_n:.3f}"],
            ["Indirect Nitrogen", f"{indirect_n:.3f}"],
            ["Methane", f"{methane:.3f}"],
            ["Supply Chain and Operational Emissions", f"{supply_chain:.3f}"],
            ["Total Greenhouse Gas Emissions Reductions", f"{total_ghg:.3f}"],
            ["Total Soil Carbon Removals", f"{soil_carbon:.3f}"],
            ["Total Impact", f"{total_impact:.3f}"]
        ]
        document.add_paragraph()
        create_table(document, impact_data, ["Impact Breakdown", "Change in mtCO2e"], font_name, font_size, font_color)

        # Water Impact Explained section
        para = document.add_paragraph()
        run = para.add_run("Water Impact Explained")
        run.font.name = font_name
        run.font.size = Pt(14)
        run.font.bold = True
        run.font.underline = True
        run.font.color.rgb = green_color
        document.add_paragraph("ESMC's water quality tool projects how much nitrogen, phosphorous and sediment in run-off water is reduced as a result of your in-field management. ESMC modeled X acres and found that your practice changes resulted in:")
        # Get water impacts for this producer's fields
        producer_fields = group['Field ID'].tolist()
        producer_water_data = water_data[water_data['id'].isin(producer_fields)]
        if not producer_water_data.empty:
            total_n_reduction = (producer_water_data['b_run_n'] - producer_water_data['p_run_n']).sum()
            total_p_reduction = (producer_water_data['b_run_p'] - producer_water_data['p_run_p']).sum()
            total_s_reduction = (producer_water_data['b_run_s'] - producer_water_data['p_run_s']).sum()
            water_modeled_acres = producer_water_data['area_ac'].sum()
            document.add_paragraph(f"ESMC modeled {water_modeled_acres:.1f} acres and found:")
            add_custom_bullet(document, "Nitrogen Reduction: ", f"{total_n_reduction:.1f} fewer pounds of nitrogen in run-off water")
            add_custom_bullet(document, "Phosphorous Reduction: ", f"{total_p_reduction:.1f} fewer pounds of phosphorous in run-off water")
            add_custom_bullet(document, "Sediment Reduction: ", f"{total_s_reduction:.1f} less tons of sediment in run-off water")
        else:
            add_custom_bullet(document, "Nitrogen Reduction: ", "X fewer pounds of nitrogen in run-off water")
            add_custom_bullet(document, "Phosphorous Reduction: ", "X fewer pounds of phosphorous in run-off water")
            add_custom_bullet(document, "Sediment Reduction: ", "X less ton of sediment in run-off water")

        # Thank you note
        document.add_paragraph()
        thankyou_para = document.add_paragraph()
        run_regular = thankyou_para.add_run("Thank you for participating in the Eco-Harvest program and for your dedication to regenerative agriculture.")
        run_regular.font.name = font_name
        run_regular.font.size = font_size

        # Contact information
        contact_para = document.add_paragraph()
        run_bold = contact_para.add_run('Need Assistance? ')
        run_bold.font.bold = True
        run_bold.font.name = font_name
        run_bold.font.size = font_size
        run_regular = contact_para.add_run("Contact your Project Representative or ")
        run_regular.font.name = font_name
        run_regular.font.size = font_size
        add_hyperlink(contact_para, 'support@ecoharvest.ag', 'http://www.ecoharvest.ag')

    return doc.producers

# Main execution
if __name__ == "__main__":
    # Load data
    carbon_data, water_data = load_data()

    # Initialize document dictionary
    class DocumentHolder:
        def __init__(self):
            self.producers = {}

    doc = DocumentHolder()

    # Generate reports
    generate_report(carbon_data, water_data, doc)

    # Save documents
    output_dir = "Producers"
    if not os.path.exists(output_dir):
        os.makedirs(output_dir)

    for producer, document in doc.producers.items():
        display_name = producer_mapping.get(producer, producer)
        try:
            safe_producer_name = producer.replace('/', '_').replace('\\', '_')
            safe_display_name = display_name.replace('/', '_').replace('\\', '_')
            doc_file_name = f"{safe_display_name} ({safe_producer_name}).docx"
        except:
            safe_producer_name = producer.replace('/', '_').replace('\\', '_')
            doc_file_name = f"{safe_producer_name}.docx"
        document.save(os.path.join(output_dir, doc_file_name))

    print("Reports generated successfully.") 

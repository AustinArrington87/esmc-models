import pandas as pd
from docx import Document
from docx.shared import Inches, Pt
from docx.oxml import OxmlElement
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.text import WD_UNDERLINE
from docx.shared import RGBColor
from docx.oxml.ns import qn
import os

# year generating reports for
run_year = "2023"
# used for mapping anonymized producer name 
mapping_df = pd.read_csv('Mapping.csv')
producer_mapping = dict(zip(mapping_df['project_producer'], mapping_df['display_name']))

def load_data():
    # Load carbon data
    carbon_data = pd.read_csv('2023_Verified.csv')
    
    # Load water impacts data
    water_data = pd.read_excel('2024_PLET_Outcomes.xlsx')
    
    return carbon_data, water_data

def format_cell_data(data):
    try:
        float_data = float(data)
        return f"{float_data:.3f}"
    except ValueError:
        return data

def create_table(doc, data, headers, font_name, font_size, font_color, scale_emissions=False):
    table = doc.add_table(rows=len(data) + 1, cols=len(headers))
    table.style = 'Table Grid'

    for i, header in enumerate(headers):
        cell = table.cell(0, i)
        run = cell.paragraphs[0].add_run(header)
        run.font.name = font_name
        run.font.size = font_size
        run.font.color.rgb = font_color
        run.font.bold = True
        cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    for row_idx, row_data in enumerate(data, start=1):
        for col_idx, cell_data in enumerate(row_data):
            if scale_emissions and isinstance(cell_data, (int, float)):
                cell_data *= 0.001
            if isinstance(cell_data, float):
                cell_data = f"{cell_data:.3f}"
            cell = table.cell(row_idx, col_idx)
            run = cell.paragraphs[0].add_run(str(cell_data))
            run.font.name = font_name
            run.font.size = font_size
            run.font.color.rgb = font_color
            cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

def calculate_delta(group, baseline_column, practice_column):
    return group[baseline_column] - group[practice_column]

def add_custom_bullet(doc, label, value):
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(label)
    run.bold = True
    run.font.name = 'Calibri'
    run.font.size = Pt(12)

    value_run = p.add_run(value)
    value_run.font.name = 'Calibri'
    value_run.font.size = Pt(12)

def add_hyperlink(paragraph, text, url):
    part = paragraph.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)

    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)
    hyperlink.set(qn('w:history'), '1')

    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    c = OxmlElement('w:color')
    c.set(qn('w:val'), '0000FF')
    rPr.append(c)
    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'single')
    rPr.append(u)
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)

    return hyperlink

def initialize_document():
    doc = Document()
    # Add header image only once
    section = doc.sections[0]
    header = section.header
    header_para = header.paragraphs[0]
    run = header_para.add_run()
    run.add_picture('header.png', width=Inches(6))
    return doc 

def generate_report(carbon_data, water_data, doc):
    font_name = 'Calibri'
    font_size = Pt(12)
    font_color = RGBColor(0, 0, 0)

    for producer, group in carbon_data.groupby('Producer (Project)'):
        if producer not in doc.producers:
            doc.producers[producer] = initialize_document()

        document = doc.producers[producer]

        # Title and Producer Info
        document.add_paragraph()
        para = document.add_heading()
        run = para.add_run("2025 Eco-Harvest Report")
        run_font = run.font
        run_font.name = 'Calibri'
        run_font.size = Pt(24)
        run_font.underline = WD_UNDERLINE.SINGLE
        run_font.color.rgb = RGBColor(0, 0, 0)

        # Producer Info
        summary_para1 = document.add_paragraph()
        run_bold = summary_para1.add_run('Producer: ')
        run_bold.font.bold = True
        run_bold.font.name = 'Calibri'
        run_bold.font.size = Pt(12)
        try:
            display_name = producer_mapping.get(producer, producer)
            run_regular = summary_para1.add_run(f"{display_name} ({producer})")
        except:
            run_regular = summary_para1.add_run(producer)
        run_regular.font.name = 'Calibri'
        run_regular.font.size = Pt(12)

        # Summary Section
        para = document.add_heading()
        run = para.add_run("Summary")
        run_font = run.font
        run_font.name = 'Calibri'
        run_font.size = Pt(16)
        run_font.underline = WD_UNDERLINE.SINGLE

        # Calculate carbon impacts
        total_reductions = group['reduced'].sum()
        total_removals = group['removed'].sum()
        total_carbon_g = total_reductions + total_removals
        total_acres = group['Acres'].sum()

        add_custom_bullet(document, "Total Carbon Impacts Generated: ", f"{total_carbon_g:.3f} metric tonnes of carbon dioxide equivalent (mtCO2e)")
        add_custom_bullet(document, "Emission Reductions: ", f"Your practice changes avoided {total_reductions:.3f} mtCO2e.")
        add_custom_bullet(document, "Carbon Removals: ", f"Your soils removed {total_removals:.3f} mtCO2e")
        add_custom_bullet(document, "Total Modeled Acres: ", f"{total_acres:.3f}")

        # Water Impacts Section
        para = document.add_heading()
        run = para.add_run("Water Impacts")
        run_font = run.font
        run_font.name = 'Calibri'
        run_font.size = Pt(16)
        run_font.underline = WD_UNDERLINE.SINGLE

        # Get water impacts for this producer's fields
        producer_fields = group['Field ID'].tolist()
        producer_water_data = water_data[water_data['id'].isin(producer_fields)]
        
        if not producer_water_data.empty:
            total_n_reduction = (producer_water_data['b_run_n'] - producer_water_data['p_run_n']).sum()
            total_p_reduction = (producer_water_data['b_run_p'] - producer_water_data['p_run_p']).sum()
            total_s_reduction = (producer_water_data['b_run_s'] - producer_water_data['p_run_s']).sum()
            water_modeled_acres = producer_water_data['area_ac'].sum()

            document.add_paragraph(f"ESMC modeled {water_modeled_acres:.1f} acres and found:")
            add_custom_bullet(document, "Nitrogen Reduction: ", f"{total_n_reduction:.1f} fewer pounds of nitrogen in run-off water")
            add_custom_bullet(document, "Phosphorous Reduction: ", f"{total_p_reduction:.1f} fewer pounds of phosphorous in run-off water")
            add_custom_bullet(document, "Sediment Reduction: ", f"{total_s_reduction:.1f} less tons of sediment in run-off water")

        # Field Results Section
        para = document.add_heading()
        run = para.add_run("Field Results")
        run_font = run.font
        run_font.name = 'Calibri'
        run_font.size = Pt(16)
        run_font.underline = WD_UNDERLINE.SINGLE

        # Carbon Results Table
        carbon_results = group[['Field Name (MRV)', 'Acres', 'reduced', 'removed']].copy()
        carbon_results['Total Impact'] = carbon_results['reduced'] + carbon_results['removed']
        carbon_results.rename(columns={
            'Field Name (MRV)': 'Field',
            'reduced': 'Emission Reductions',
            'removed': 'Carbon Removals'
        }, inplace=True)
        
        create_table(document, carbon_results.values, 
                    ['Field', 'Acres', 'Emission Reductions', 'Carbon Removals', 'Total Impact'],
                    font_name, font_size, font_color)

        # Water Results Table (if available)
        if not producer_water_data.empty:
            document.add_paragraph()
            para = document.add_heading()
            run = para.add_run("Water Impact Results")
            run_font = run.font
            run_font.name = 'Calibri'
            run_font.size = Pt(14)

            water_results = producer_water_data[['field_id', 'area_ac', 
                                               'b_run_n', 'p_run_n', 
                                               'b_run_p', 'p_run_p',
                                               'b_run_s', 'p_run_s']].copy()
            
            water_results['N Reduction'] = water_results['b_run_n'] - water_results['p_run_n']
            water_results['P Reduction'] = water_results['b_run_p'] - water_results['p_run_p']
            water_results['S Reduction'] = water_results['b_run_s'] - water_results['p_run_s']
            
            water_results = water_results[['field_id', 'area_ac', 'N Reduction', 'P Reduction', 'S Reduction']]
            water_results.rename(columns={
                'field_id': 'Field',
                'area_ac': 'Acres',
                'N Reduction': 'Nitrogen Reduction (lbs)',
                'P Reduction': 'Phosphorous Reduction (lbs)',
                'S Reduction': 'Sediment Reduction (tons)'
            }, inplace=True)

            create_table(document, water_results.values,
                        ['Field', 'Acres', 'Nitrogen Reduction (lbs)', 
                         'Phosphorous Reduction (lbs)', 'Sediment Reduction (tons)'],
                        font_name, font_size, font_color)

        # Thank you note
        document.add_paragraph()
        thankyou_para = document.add_paragraph()
        run_regular = thankyou_para.add_run("Thank you for participating in the Eco-Harvest program and for your dedication to regenerative agriculture.")
        run_regular.font.name = 'Calibri'
        run_regular.font.size = Pt(12)

        # Contact information
        contact_para = document.add_paragraph()
        run_bold = contact_para.add_run('Need Assistance? ')
        run_bold.font.bold = True
        run_bold.font.name = 'Calibri'
        run_bold.font.size = Pt(12)
        run_regular = contact_para.add_run("Contact your Project Representative or ")
        run_regular.font.name = 'Calibri'
        run_regular.font.size = Pt(12)
        add_hyperlink(contact_para, 'support@ecoharvest.ag', 'http://www.ecoharvest.ag')

    return doc.producers

# Main execution
if __name__ == "__main__":
    # Load data
    carbon_data, water_data = load_data()

    # Initialize document dictionary
    class DocumentHolder:
        def __init__(self):
            self.producers = {}

    doc = DocumentHolder()

    # Generate reports
    generate_report(carbon_data, water_data, doc)

    # Save documents
    output_dir = "Producers"
    if not os.path.exists(output_dir):
        os.makedirs(output_dir)

    for producer, document in doc.producers.items():
        display_name = producer_mapping.get(producer, producer)
        try:
            safe_producer_name = producer.replace('/', '_').replace('\\', '_')
            safe_display_name = display_name.replace('/', '_').replace('\\', '_')
            doc_file_name = f"{safe_display_name} ({safe_producer_name}).docx"
        except:
            safe_producer_name = producer.replace('/', '_').replace('\\', '_')
            doc_file_name = f"{safe_producer_name}.docx"
        document.save(os.path.join(output_dir, doc_file_name))

    print("Reports generated successfully.") 

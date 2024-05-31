import pandas as pd
from docx import Document
from docx.shared import Inches, Pt
from docx.oxml import OxmlElement
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.text import WD_UNDERLINE
from docx.shared import RGBColor
from docx.oxml.ns import qn
import os

# Load the data
year = "2023"
batch = "Verified"
file_path = year+'_'+batch+'_Adjusted.csv'
data = pd.read_csv(file_path)

# Define the standard style settings
font_name = 'Calibri'
font_size = Pt(12)
font_color = RGBColor(0, 0, 0)

# Create the 'Producers' directory if it doesn't exist
output_dir = "Producers"
if not os.path.exists(output_dir):
    os.makedirs(output_dir)

def format_cell_data(data):
    # Format data to 3 decimal places if it is a number
    try:
        # Check if the data is convertible to a float
        float_data = float(data)
        # Return formatted data
        return f"{float_data:.3f}"
    except ValueError:
        # Return original data if it's not a number
        return data


# functon for creating summary tables T

def create_table(doc, data, headers, font_name, font_size, font_color, scale_emissions=False):
    # Create a table with the specified data and headers
    table = doc.add_table(rows=len(data) + 1, cols=len(headers))
    table.style = 'Table Grid'

    # Insert column headers and format them
    for i, header in enumerate(headers):
        cell = table.cell(0, i)
        run = cell.paragraphs[0].add_run(header)
        run.font.name = font_name
        run.font.size = font_size
        run.font.color.rgb = font_color
        run.font.bold = True  # Making headers bold
        cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Insert the data into the table
    for row_idx, row_data in enumerate(data, start=1):
        for col_idx, cell_data in enumerate(row_data):
            if scale_emissions and isinstance(cell_data, (int, float)):
                cell_data *= 0.001  # Scale the data if required
            
            if isinstance(cell_data, float):
                cell_data = f"{cell_data:.3f}"  # Truncate to 3 decimal places
            cell = table.cell(row_idx, col_idx)
            run = cell.paragraphs[0].add_run(str(cell_data))
            run.font.name = font_name
            run.font.size = font_size
            run.font.color.rgb = font_color
            cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # Aligning the cell content to center


# Calculate Delta for N2O, Methane, Fossil Fuel, Pesticides, and Upstream Emissions
def calculate_delta(group, baseline_column, practice_column):
    return group[baseline_column] - group[practice_column]

# custom font for bulleted sections 
def add_custom_bullet(doc, label, value):
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(label)
    run.bold = True
    run.font.name = 'Calibri'
    run.font.size = Pt(12)  # Setting the font size to 11, modify as needed

    value_run = p.add_run(value)
    value_run.font.name = 'Calibri'
    value_run.font.size = Pt(12)

def add_custom_bullet_link(doc, label, value, hyperlink_text, hyperlink_url):
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(label)
    run.bold = True
    run.font.name = 'Calibri'
    run.font.size = Pt(12)

    # Add initial text before hyperlink
    value_parts = value.split(hyperlink_text)
    if len(value_parts) > 1:
        value_run = p.add_run(value_parts[0])
        value_run.font.name = 'Calibri'
        value_run.font.size = Pt(12)

        # Add hyperlink text
        add_hyperlink(p, hyperlink_text, hyperlink_url)

        # Add text after hyperlink
        value_run = p.add_run(value_parts[1])
        value_run.font.name = 'Calibri'
        value_run.font.size = Pt(12)
    else:
        value_run = p.add_run(value)
        value_run.font.name = 'Calibri'
        value_run.font.size = Pt(12)

    return p


# Hyperlink 

def add_hyperlink(paragraph, text, url):
    """
    A function that places a hyperlink within a paragraph object.

    :param paragraph: The paragraph we are adding the hyperlink to.
    :param text: The text to be used for the link.
    :param url: The URL the link points to.
    """
    # This gets access to the document.xml.rels file and gets a new relation id value
    part = paragraph.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)

    # Create the w:hyperlink tag and set needed values
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id, )
    hyperlink.set(qn('w:history'), '1')  # optional

    # Create a w:r element
    new_run = OxmlElement('w:r')

    # Create a new w:rPr element
    rPr = OxmlElement('w:rPr')

    # Add color if needed
    c = OxmlElement('w:color')
    c.set(qn('w:val'), '0000FF')  # sets the color to blue
    rPr.append(c)

    # Add underline if needed
    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'single')
    rPr.append(u)

    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)

    paragraph._p.append(hyperlink)

    return hyperlink

######### MAIN FOR LOOP to Generate Reports ##########
# Generate a report for each producer
for producer, group in data.groupby('Producer (Project)'):
    doc = Document()

    # Access the 'Project' column of the first row in the group
    project = group['Project'].iloc[0]

    # Access the header of the document
    section = doc.sections[0]

    header = section.header
    # Add a paragraph in the header and then add a run
    header_para = header.paragraphs[0]
    run = header_para.add_run()
    # Add your PNG image to the run (specify the path to your image)
    run.add_picture('header.png', width=Inches(11))
    doc.add_paragraph() # empty space 

    # producer name 
    #doc.add_heading(producer, 0)

    para = doc.add_heading()
    run = para.add_run(f"Eco-Harvest Carbon Impact Report")
    run_font = run.font
    run_font.name = 'Calibri'
    run_font.size = Pt(24)
    run_font.underline = WD_UNDERLINE.SINGLE
    run_font.color.rgb = RGBColor(0, 0, 0)

    # summary info at top 
    summary_para1 = doc.add_paragraph()
    # Add a run for "Producer:" and make it bold
    run_bold = summary_para1.add_run('Producer: ')
    run_bold.font.bold = True
    run_bold.font.name = 'Calibri'
    run_bold.font.size = Pt(12)
    run_bold.font.color.rgb = RGBColor(0, 0, 0)

    # Add another run for "Producer-1" without bold
    run_regular = summary_para1.add_run(producer)
    run_regular.font.bold = False  # Not necessary as False is the default
    run_regular.font.name = 'Calibri'
    run_regular.font.size = Pt(12)
    run_regular.font.color.rgb = RGBColor(0, 0, 0)

    # Add a run for "Project:" and make it bold
    summary_para2 = doc.add_paragraph()
    run_bold = summary_para2.add_run('Project: ')
    run_bold.font.bold = True
    run_bold.font.name = 'Calibri'
    run_bold.font.size = Pt(12)
    run_bold.font.color.rgb = RGBColor(0, 0, 0)

    # Add another run for "Project" without bold
    run_regular = summary_para2.add_run(project)
    run_regular.font.bold = False  # Not necessary as False is the default
    run_regular.font.name = 'Calibri'
    run_regular.font.size = Pt(12)
    run_regular.font.color.rgb = RGBColor(0, 0, 0)

    # Add a run for "Year" and make it bold
    summary_para3 = doc.add_paragraph()
    run_bold = summary_para3.add_run('Year: ')
    run_bold.font.bold = True
    run_bold.font.name = 'Calibri'
    run_bold.font.size = Pt(12)
    run_bold.font.color.rgb = RGBColor(0, 0, 0)

    # Add another run for "Year" without bold
    run_regular = summary_para3.add_run(year)
    run_regular.font.bold = False  # Not necessary as False is the default
    run_regular.font.name = 'Calibri'
    run_regular.font.size = Pt(12)
    run_regular.font.color.rgb = RGBColor(0, 0, 0)

    # Add a run for "Batch" and make it bold
    summary_para4 = doc.add_paragraph()
    run_bold = summary_para4.add_run('Batch: ')
    run_bold.font.bold = True
    run_bold.font.name = 'Calibri'
    run_bold.font.size = Pt(12)
    run_bold.font.color.rgb = RGBColor(0, 0, 0)

    # Add another run for "Batch" without bold
    run_regular = summary_para4.add_run(batch)
    run_regular.font.bold = False  # Not necessary as False is the default
    run_regular.font.name = 'Calibri'
    run_regular.font.size = Pt(12)
    run_regular.font.color.rgb = RGBColor(0, 0, 0)


    # YOUR RESULTS 
    para = doc.add_heading()
    run = para.add_run(f"Your Results")
    run_font = run.font
    run_font.name = 'Calibri'
    run_font.size = Pt(16)
    run_font.underline = WD_UNDERLINE.SINGLE
    run_font.color.rgb = RGBColor(0, 0, 0)

    # Fields Summary section load 
    fields_data = group[['Field Name (MRV)', 'Acres', 'Commodity']].copy()  # Replace 'Commodity'
    fields_data.rename(columns={'Field Name (MRV)': 'Field'}, inplace=True)
    # Reductions section load 
    #reductions_data = group[['Field Name (MRV)', 'reduced']].copy()
    reductions_data = group[['Field Name (MRV)', 'reduced_adjusted']].copy()
    reductions_data.rename(columns={'Field Name (MRV)': 'Field'}, inplace=True)
    total_reductions = reductions_data['reduced_adjusted'].sum()
    # Removals section load 
    removals_data = group[['Field Name (MRV)', 'baseline_dsoc', 'dsoc', 'removed']].copy()
    removals_data.rename(columns={'Field Name (MRV)': 'Field', 'removed': 'Removed'}, inplace=True)

    total_removals = removals_data['Removed'].sum()
    total_carbon_g = total_reductions + total_removals

    # Total Carbon Impacts
    add_custom_bullet(doc, "Total Carbon Impacts Generated: ", f"{total_carbon_g:.3f} metric tonnes of carbon dioxide equivalent (mtCO2e)")

    # Create a bullet point for Total Reductions and apply styles
    add_custom_bullet(doc, "Emission Reductions: ", f"Your practice changes avoided {total_reductions:.3f} mtCO2e.")

    # Create a bullet point for Total Removals and apply styles
    add_custom_bullet(doc, "Carbon Removals: ", f"Your soils removed {total_removals:.3f} mtCO2e")

    # Create a bullet point for Total Acres and apply styles
    add_custom_bullet(doc, "Total Modeled Acres*: ", f"{fields_data['Acres'].sum():.3f}")


    # Add another run for note aboute modeled acres
    summary_para3 = doc.add_paragraph()
    run_regular = summary_para3.add_run("*ESMC requires complete historical data to generate a baseline, so fields with historical data gaps are not modeled.")
    run_regular.font.bold = False  # Not necessary as False is the default
    run_regular.font.name = 'Calibri'
    run_regular.font.size = Pt(9)
    run_regular.font.color.rgb = RGBColor(0, 0, 0)

    # Eco-Harvest Program Info 
    para = doc.add_heading()
    run = para.add_run(f"Eco-Harvest: rewarding producers for regenerative ag ")
    run_font = run.font
    run_font.name = 'Calibri'
    run_font.size = Pt(16)
    run_font.underline = WD_UNDERLINE.SINGLE
    run_font.color.rgb = RGBColor(0, 0, 0)

    add_custom_bullet(doc, "Customizable practices: ", f"Choose how to implement regenerative practices.")
    add_custom_bullet(doc, "Fair Payments: ", f"As a non-profit, ESMC maximizes producer benefits.")
    add_custom_bullet_link(doc, "Return on Investment: ", 
                  "Many farmers experience a 15-25% return on investment after 3–5 years of regenerative farming.", 
                  "regenerative farming", 
                  "https://www.bcg.com/publications/2023/regenerative-agriculture-benefits-germany-beyond")
    add_custom_bullet(doc, "Weather-Ready Farming: ", f"Thank you for participating in the Eco-Harvest program. Continue working with your advisor to improve your soil health and boost your farm’s productivity and earnings.")

    # Thank You 
    thanks_para1 = doc.add_paragraph()
    # Add another run for "Producer-1" without bold
    run_regular = thanks_para1.add_run("Thank you for participating in the Eco-Harvest program. Continue working with your advisor to improve your soil health and boost your farm’s productivity and earnings.")
    run_regular.font.name = 'Calibri'
    run_regular.font.size = Pt(12)
    run_regular.font.color.rgb = RGBColor(0, 0, 0)

    # Contact Us 
     # summary info at top 
    summary_para1 = doc.add_paragraph()
    # Add a run for "Producer:" and make it bold
    run_bold = summary_para1.add_run('Need Assistance? ')
    run_bold.font.bold = True
    run_bold.font.name = 'Calibri'
    run_bold.font.size = Pt(12)
    run_bold.font.color.rgb = RGBColor(0, 0, 0)

    # Add another run for "Producer-1" without bold
    run_regular = summary_para1.add_run("Contact your Project Representative or ")
    run_regular.font.bold = False  # Not necessary as False is the default
    run_regular.font.name = 'Calibri'
    run_regular.font.size = Pt(12)
    run_regular.font.color.rgb = RGBColor(0, 0, 0)
    add_hyperlink(summary_para1, 'support@ecoharvest.ag', 'http://www.ecoharvest.ag')


    ############## SECTION 2 ##################################


    ## Soil Data section
    para = doc.add_heading()
    run = para.add_run(f"Soil Sampling Results")
    run_font = run.font
    run_font.name = 'Calibri'
    run_font.size = Pt(14)
    run_font.underline = WD_UNDERLINE.SINGLE
    run_font.color.rgb = RGBColor(112, 173, 71)

    # Thank You 
    thanks_para2 = doc.add_paragraph()
    # Add another run for "Producer-1" without bold
    run_regular = thanks_para2.add_run("Thank you for participating in the Eco-Harvest program. Continue working with your advisor to improve your soil health and boost your farm’s productivity and earnings.")
    run_regular.font.name = 'Calibri'
    run_regular.font.size = Pt(12)
    run_regular.font.color.rgb = RGBColor(0, 0, 0)


    soil_bullet_points = [
    "SOC – The percentage of soil that is organic carbon.  2-6% is optimal.",
    "pH – A measurement of acidity or alkalinity of the soil. Optimal rankings are 4-8.",
    "BD – Bulk Density is the oven dry weight of soil per unit volume (g/cm3) ",
    "Clay - The percentage of your soil that is clay, 20-35% is optimal. "
    ]
    for point in soil_bullet_points:
        paragraph = doc.add_paragraph()
        run_bullet_regular = paragraph.add_run(f'• {point}')
        run_bullet_regular.font.name = font_name
        run_bullet_regular.font.size = font_size
        run_bullet_regular.font.color.rgb = font_color

    ## Soil Table 
    soil_data = group[['Field Name (MRV)', 'soil_avg_soc', 'soil_avg_ph', 'soil_avg_bulkdensity', 'soil_clay_fraction']].copy()
    soil_data.rename(columns={'Field Name (MRV)': 'Field', 'soil_avg_soc': 'SOC', 'soil_avg_ph': 'pH', 'soil_avg_bulkdensity': 'BD', 'soil_clay_fraction': 'Clay'}, inplace=True)
    #create_table(doc, soil_data.values, ['Field', 'SOC', 'pH', 'BD', 'Clay'])
    # Create a table with the specified data and headers
    headers = ['Field', 'SOC', 'pH', 'BD', 'Clay']
    data = soil_data.values
    table = doc.add_table(rows=len(data) + 1, cols=len(headers))
    table.style = 'Table Grid'

    # Insert column headers and format them
    for i, header in enumerate(headers):
        cell = table.cell(0, i)
        run = cell.paragraphs[0].add_run(header)
        run.font.name = font_name
        run.font.size = font_size
        run.font.color.rgb = font_color
        run.font.bold = True  # Making headers bold
        cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Insert the data into the table
    for row_idx, row_data in enumerate(data, start=1):
        for col_idx, cell_data in enumerate(row_data):
            cell = table.cell(row_idx, col_idx)
            if isinstance(cell_data, float):
                cell_data = f"{cell_data:.3f}"
            run = cell.paragraphs[0].add_run(str(cell_data))
            run.font.name = font_name
            run.font.size = font_size
            run.font.color.rgb = font_color

    ####################################

    ## Carbon Impact section
    para = doc.add_heading()
    run = para.add_run(f"Carbon Impact Explained")
    run_font = run.font
    run_font.name = 'Calibri'
    run_font.size = Pt(14)
    run_font.underline = WD_UNDERLINE.SINGLE
    run_font.color.rgb = RGBColor(112, 173, 71)

    # Thank You - Carbon 
    thanks_para3 = doc.add_paragraph()
    # Add another run for "Producer-1" without bold
    run_regular = thanks_para3.add_run("ESMC uses historical data, a scientific model, soil samples and weather data to calculate emissions reductions and soil carbon removals. Any negative emissions are increased emissions, which is not unusual in the early years of adopting a practice change. If concerned, please speak with your technical advisor about further modifications to your practice change(s).")
    run_regular.font.name = 'Calibri'
    run_regular.font.size = Pt(12)
    run_regular.font.color.rgb = RGBColor(0, 0, 0)


     # Carbon Info 

    add_custom_bullet(doc, "Greenhouse Gas Emissions (GHG) Reductions (mtCO2e)  ", f" are the sum of the reductions of methane (CH4), direct nitrous oxide (N2O), indirect nitrous oxide (N2O) and supply chain and operational emissions** between the baseline year and the project year.")
    add_custom_bullet(doc, "Soil Carbon Removals (mtCO2e): ", f" The difference in the amount of carbon in the soil between the baseline year and the project year. ")


     # Add another run for note aboute modeled acres
    summary_para3 = doc.add_paragraph()
    run_regular = summary_para3.add_run("**Upstream and on-field process emissions associated with management such as fuel used on farm and imbedded emissions of fertilizer products; ESMC utilized emission factors from Ecoinvent and WFLD.")
    run_regular.font.bold = False  # Not necessary as False is the default
    run_regular.font.name = 'Calibri'
    run_regular.font.size = Pt(9)
    run_regular.font.color.rgb = RGBColor(0, 0, 0)

     
    # IMPACT Breakdown Table 
    

    ############## SECTION 3 ##################################


    # Similar approach for N2O, Methane, etc.
    
    # Direct N2O
    group['Delta Direct N2O'] = calculate_delta(group, 'baseline_n20_direct', 'n2o_direct')
    # N2O Direct section
    n2o_data = group[['Field Name (MRV)', 'baseline_n20_direct', 'n2o_direct', 'Delta Direct N2O']].copy()
    n2o_data.rename(columns={
        'Field Name (MRV)': 'Field',
        'baseline_n20_direct': 'Direct N2O Baseline',
        'n2o_direct': 'Direct N2O Practice',
        'Delta Direct N2O': 'Delta'
    }, inplace=True)
 
    total_delta_direct_n2o = n2o_data['Delta'].sum()

    #Indirect N2O
    group['Delta Indirect N2O'] = calculate_delta(group, 'baseline_n2o_indirect_adjusted', 'n2o_indirect_adjusted')
    # N2O Direct section
    n2oindirect_data = group[['Field Name (MRV)', 'baseline_n2o_indirect_adjusted', 'n2o_indirect_adjusted', 'Delta Indirect N2O']].copy()
    n2oindirect_data.rename(columns={
        'Field Name (MRV)': 'Field',
        'baseline_n2o_indirect_adjusted': 'Indirect N2O Baseline',
        'n2o_indirect_adjusted': 'Indirect N2O Practice',
        'Delta Indirect N2O': 'Delta'
    }, inplace=True)

    total_delta_indirect_n2o = n2oindirect_data['Delta'].sum()
 
    # Methane 
    group['Delta Methane'] = calculate_delta(group, 'baseline_methane_adjusted', 'methane_adjusted')
    # Methane section
    methane_data = group[['Field Name (MRV)', 'baseline_methane_adjusted', 'methane_adjusted', 'Delta Methane']].copy()
    methane_data.rename(columns={
        'Field Name (MRV)': 'Field',
        'baseline_methane_adjusted': 'CH4 Baseline',
        'methane_adjusted': 'CH4 Practice',
        'Delta Methane': 'Delta'
    }, inplace=True)
 
    total_delta_methane = methane_data['Delta'].sum()
  
    #Emission Factors 
    group['Delta Field Emissions'] = calculate_delta(group, 'field_baseline_emissions', 'field_practice_emissions')
    # Emissions section
    emissions_data = group[['Field Name (MRV)', 'field_baseline_emissions', 'field_practice_emissions', 'Delta Field Emissions']].copy()
    emissions_data.rename(columns={
        'Field Name (MRV)': 'Field',
        'field_baseline_emissions': 'Supply Chain & Operational Emissions Baseline',
        'field_practice_emissions': 'Supply Chain & Operational Emissions Practice',
        'Delta Field Emissions': 'Delta'
    }, inplace=True)

    # Summarize the total Emissions
    total_delta_emissions = emissions_data['Delta'].sum()
    # convert kg to CO23
    total_delta_emissions = total_delta_emissions * 0.001
 
    ### IMPACT TABLE 
    table = doc.add_table(rows=8, cols=2)
    table.style = 'Table Grid'


    # Define the headers
    headers = ['Impact Breakdown', 'Change between Baseline and practice years (mtCO2e)']

    # Insert column headers and format them
    for i, header in enumerate(headers):
        cell = table.cell(0, i)
        cell.text = header
        run = cell.paragraphs[0].runs[0]
        run.font.name = font_name
        run.font.size = font_size
        run.font.color.rgb = font_color
        if i == 0:
            run.font.bold = True
        else:
            run.font.italic = True
        cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER


    # Define the rows under "Impact Breakdown"
    impact_breakdown = [
        "Greenhouse Gas Emissions Reductions",
        "       N2O Indirect",
        "       N2O Direct",
        "       Methane",
        "       Supply Chain and Operational Emissions",
        "Total Soil Carbon Removals",
        "Total Carbon Impact (Reductions + Removals)"
        ]
    # Values to be bolded in the Impact Breakdown column
    bold_values = [
        "Greenhouse Gas Emissions Reductions",
        "Total Soil Carbon Removals",
        "Total Carbon Impact (Reductions + Removals)"
        ]

    total_carbon_impact = total_reductions + total_removals

    # List of corresponding values for "DELTA mtCO2e"

    delta_values = [
    round(total_reductions, 3),
    round(total_delta_indirect_n2o, 3),
    round(total_delta_direct_n2o, 3),
    round(total_delta_methane, 3),
    round(total_delta_emissions, 3),
    round(total_removals, 3),
    round(total_carbon_impact, 3)
    ]

    # Populate the table rows with data
    for row in range(1, 8):  # Start from 1 to leave the header row as is
        impact_cell = table.cell(row, 0)
        delta_cell = table.cell(row, 1)

        impact_run = impact_cell.paragraphs[0].add_run(impact_breakdown[row - 1])
        delta_run = delta_cell.paragraphs[0].add_run(str(delta_values[row - 1]))

        # Apply standard font settings to impact breakdown cell
        impact_run.font.name = font_name
        impact_run.font.size = font_size
        impact_run.font.color.rgb = font_color

        # Apply standard font settings to delta value cell
        delta_run.font.name = font_name
        delta_run.font.size = font_size
        delta_run.font.color.rgb = font_color

        # Check if the impact breakdown should be bold
        if impact_breakdown[row - 1] in bold_values:
            impact_run.font.bold = True

    ## Fields 
    para = doc.add_heading()
    run = para.add_run(f"Fields Results")
    run_font = run.font
    run_font.name = 'Calibri'
    run_font.size = Pt(14)
    run_font.underline = WD_UNDERLINE.SINGLE
    run_font.color.rgb = RGBColor(112, 173, 71)

    # REDUCTIONS 
    para = doc.add_heading()
    run = para.add_run(f"Reductions")
    run_font = run.font
    run_font.name = 'Calibri'
    run_font.size = Pt(12)

    doc.add_paragraph("This section shows the GHG emissions reductions from the intervention(s) per field. Total reductions include the following emissions: direct nitrogen, methane and supply chain & operational emissions. ")

    doc.add_paragraph()
    total_reductions = reductions_data['reduced_adjusted'].sum()
    para = doc.add_paragraph()
    run = para.add_run(f"Total Reductions: {total_reductions:.3f} tonnes CO2e")
    run.italic = True
    
    # Example for Total Reductions
    reductions_data = group[['Field Name (MRV)', 'reduced_adjusted']].copy()
    reductions_data.rename(columns={'Field Name (MRV)': 'Field', 'reduced_adjusted': 'Total Reductions'}, inplace=True)
    create_table(doc, reductions_data.values, ['Field', 'Total Reductions'], font_name, font_size, font_color)

    # Direct Nitrogen (N2O) Emissions 
    doc.add_paragraph()
    para = doc.add_paragraph()
    run = para.add_run("Direct Nitrogen (N2O) Emissions")
    run.bold = True
    run.font.name = 'Calibri'

    doc.add_paragraph()
    total_n2o_reductions = n2o_data['Delta'].sum()
    para = doc.add_paragraph()
    run = para.add_run(f"Total Direct N2O Reductions: {total_n2o_reductions:.3f} tonnes CO2e")
    run.italic = True

    # Direct n2o Table 
    create_table(doc, n2o_data.values, ['Field', 'Direct N2O Baseline', 'Direct N2O Practice', 'Delta'], font_name, font_size, font_color)


    # Indirect N2O 

    doc.add_paragraph()
    para = doc.add_paragraph()
    run = para.add_run("Indirect Nitrogen (N2O) Emissions")
    run.bold = True
    run.font.name = 'Calibri'

    doc.add_paragraph()
    total_indirectn2o_reductions = n2oindirect_data['Delta'].sum()
    para = doc.add_paragraph()
    run = para.add_run(f"Total Indirect N2O Reductions: {total_indirectn2o_reductions:.3f} tonnes CO2e")
    run.italic = True

    # Indirect n2o Table 
    create_table(doc, n2oindirect_data.values, ['Field', 'Indirect N2O Baseline', 'Indirect N2O Practice', 'Delta'], font_name, font_size, font_color)

    # Methane 

    doc.add_paragraph()
    para = doc.add_paragraph()
    run = para.add_run("Methane (CH4) Emissions")
    run.bold = True
    run.font.name = 'Calibri'

    doc.add_paragraph()
    total_ch4_reductions = methane_data['Delta'].sum()
    para = doc.add_paragraph()
    run = para.add_run(f"Total CH4 Reductions: {total_ch4_reductions:.3f} tonnes CO2e")
    run.italic = True

    # methane table 
    create_table(doc, methane_data.values, ['Field', 'CH4 Baseline', 'CH4 Practice', 'Delta'], font_name, font_size, font_color)

    # Save the document
    safe_producer_name = producer.replace('/', '_').replace('\\', '_')
    doc_file_name = f"{safe_producer_name}.docx"
    doc.save(os.path.join(output_dir, doc_file_name))

    # Supply chain 

    doc.add_paragraph()
    para = doc.add_paragraph()
    run = para.add_run("Supply Chain and Operational Emissions")
    run.bold = True
    run.font.name = 'Calibri'

    doc.add_paragraph()
    total_emissions_reductions = emissions_data['Delta'].sum()
    para = doc.add_paragraph()
    run = para.add_run(f"Total Supply Chain and Operational Emissions: {total_delta_emissions:.3f} tonnes CO2e")
    run.italic = True

    # emissions table 
    create_table(doc, emissions_data.values, ['Field', 'Supply Chain & Operational Emissions Baseline', 'Supply Chain & Operational Emissions Practice', 'Delta'], font_name, font_size, font_color, scale_emissions=True)

    # REMOVALS
    para = doc.add_heading()
    run = para.add_run(f"Removals")
    run_font = run.font
    run_font.name = 'Calibri'
    run_font.size = Pt(12)

    doc.add_paragraph("This section shows the modeled carbon removals per field from the intervention(s).  ")

    doc.add_paragraph()
    para = doc.add_paragraph()
    run = para.add_run(f"Total Reductions: {total_removals:.3f} tonnes CO2e")
    run.italic = True

    # removals table

    removals_data.rename(columns={
        'baseline_dsoc': 'Removals Baseline',
        'dsoc': 'Removals Practice',
        'Removed': 'Delta'
        }, inplace=True)

    create_table(doc, removals_data.values, ['Field', 'Removals Baseline', 'Removals Practice', 'Delta'], font_name, font_size, font_color)

    # PAYMENT 
    para = doc.add_heading()
    run = para.add_run(f"Payment Structure")
    run_font = run.font
    run_font.name = 'Calibri'
    run_font.size = Pt(12)

    doc.add_paragraph("Producers in this project received a payment of $X/acre OR $X/impact. ")

    doc.add_picture('diagram.png')
    italic_paragraph = doc.add_paragraph()
    italic_run = italic_paragraph.add_run('Figure 1: Total impact is the combined value of GHG emissions reductions and soil carbon stored. This diagram illustrates a general improvement from the baseline to project year – GHG emissions decrease and soil carbon increases.')
    italic_run.italic = True
    italic_run.font_name = 'Calibri'
    italic_run.font.size = Pt(11)

    doc.add_paragraph()
    para = doc.add_paragraph()
    run = para.add_run("Real-world Comparisons")
    run.bold = True
    run.font.name = 'Calibri'

    # Add another run for "Producer-1" without bold
    summary_para5 = doc.add_paragraph()
    run_regular = summary_para5.add_run("Use ")
    run_regular.font.bold = False  # Not necessary as False is the default
    run_regular.font.name = 'Calibri'
    run_regular.font.size = Pt(12)
    run_regular.font.color.rgb = RGBColor(0, 0, 0)
    add_hyperlink(summary_para5, 'this tool', 'https://www.epa.gov/energy/greenhouse-gas-equivalencies-calculator#resultsg')
    run_regular = summary_para5.add_run(" to compare carbon impacts to everyday examples.")
    run_regular.font.bold = False  # Not necessary as False is the default
    run_regular.font.name = 'Calibri'
    run_regular.font.size = Pt(12)

    # Save the document
    safe_producer_name = producer.replace('/', '_').replace('\\', '_')
    doc_file_name = f"{safe_producer_name}.docx"
    doc.save(os.path.join(output_dir, doc_file_name))


print("Reports generated successfully.")

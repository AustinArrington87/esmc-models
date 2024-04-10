import pandas as pd
from docx import Document
from docx.shared import Inches, Pt
from docx.oxml import OxmlElement
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.text import WD_UNDERLINE
from docx.shared import RGBColor
import os

# Load the data
year = "2021"
file_path = year+'_Verified.csv'
data = pd.read_csv(file_path)

# Create the 'Producers' directory if it doesn't exist
output_dir = "Producers"
if not os.path.exists(output_dir):
    os.makedirs(output_dir)

def format_cell_data(data):
    # Format data to 3 decimal places if it is a number
    try:
        # Check if the data is convertible to a float
        float_data = float(data)
        # Return formatted data
        return f"{float_data:.3f}"
    except ValueError:
        # Return original data if it's not a number
        return data


# Function to create and format a table in the Word document
def create_table(document, data, column_names, title=None, note=None):
    if title:
        document.add_heading(title, level=1)
    table = document.add_table(rows=1, cols=len(column_names))
    for i, column in enumerate(column_names):
        table.cell(0, i).text = column
        table.cell(0, i).paragraphs[0].runs[0].font.bold = True  # Make header bold
    for row_data in data:
        row_cells = table.add_row().cells
        for i, cell_data in enumerate(row_data):
            row_cells[i].text = format_cell_data(cell_data)
    if note:
        para = document.add_paragraph()
        run = para.add_run(note)
        run.italic = True

# Calculate Delta for N2O, Methane, Fossil Fuel, Pesticides, and Upstream Emissions
def calculate_delta(group, baseline_column, practice_column):
    return group[baseline_column] - group[practice_column]

# custom font for bulleted sections 
def add_custom_bullet(doc, label, value):
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(label)
    run.bold = True
    run.font.name = 'Calibri'
    run.font.size = Pt(11)  # Setting the font size to 11, modify as needed

    value_run = p.add_run(value)
    value_run.font.name = 'Calibri'
    value_run.font.size = Pt(11)

# Generate a report for each producer
for producer, group in data.groupby('Producer (Project)'):
    doc = Document()

    # Access the 'Project' column of the first row in the group
    project = group['Project'].iloc[0]

    # Access the header of the document
    section = doc.sections[0]

    header = section.header
    # Add a paragraph in the header and then add a run
    header_para = header.paragraphs[0]
    run = header_para.add_run()
    # Add your PNG image to the run (specify the path to your image)
    run.add_picture('header.png', width=Inches(11))
    doc.add_paragraph() # empty space 

    # producer name 
    #doc.add_heading(producer, 0)

    para = doc.add_heading()
    run = para.add_run(f"Eco-Harvest Carbon Impact Report")
    run_font = run.font
    run_font.name = 'Calibri'
    run_font.size = Pt(24)
    run_font.underline = WD_UNDERLINE.SINGLE
    run_font.color.rgb = RGBColor(0, 0, 0)

    # summary info at top 
    summary_para1 = doc.add_paragraph()
    # Add a run for "Producer:" and make it bold
    run_bold = summary_para1.add_run('Producer: ')
    run_bold.font.bold = True
    run_bold.font.name = 'Calibri'
    run_bold.font.size = Pt(12)
    run_bold.font.color.rgb = RGBColor(0, 0, 0)

    # Add another run for "Producer-1" without bold
    run_regular = summary_para1.add_run(producer)
    run_regular.font.bold = False  # Not necessary as False is the default
    run_regular.font.name = 'Calibri'
    run_regular.font.size = Pt(12)
    run_regular.font.color.rgb = RGBColor(0, 0, 0)

    # Add a run for "Project:" and make it bold
    summary_para2 = doc.add_paragraph()
    run_bold = summary_para2.add_run('Project: ')
    run_bold.font.bold = True
    run_bold.font.name = 'Calibri'
    run_bold.font.size = Pt(12)
    run_bold.font.color.rgb = RGBColor(0, 0, 0)

    # Add another run for "Project" without bold
    run_regular = summary_para2.add_run(project)
    run_regular.font.bold = False  # Not necessary as False is the default
    run_regular.font.name = 'Calibri'
    run_regular.font.size = Pt(12)
    run_regular.font.color.rgb = RGBColor(0, 0, 0)

    # Add a run for "Year" and make it bold
    summary_para3 = doc.add_paragraph()
    run_bold = summary_para3.add_run('Year: ')
    run_bold.font.bold = True
    run_bold.font.name = 'Calibri'
    run_bold.font.size = Pt(12)
    run_bold.font.color.rgb = RGBColor(0, 0, 0)

    # Add another run for "Project" without bold
    run_regular = summary_para3.add_run(year)
    run_regular.font.bold = False  # Not necessary as False is the default
    run_regular.font.name = 'Calibri'
    run_regular.font.size = Pt(12)
    run_regular.font.color.rgb = RGBColor(0, 0, 0)

    # YOUR RESULTS 
    para = doc.add_heading()
    run = para.add_run(f"Your Results")
    run_font = run.font
    run_font.name = 'Calibri'
    run_font.size = Pt(16)
    run_font.underline = WD_UNDERLINE.SINGLE
    run_font.color.rgb = RGBColor(0, 0, 0)

    # Fields section load 
    fields_data = group[['Field Name (MRV)', 'Acres', 'Commodity']].copy()  # Replace 'Commodity'
    fields_data.rename(columns={'Field Name (MRV)': 'Field'}, inplace=True)
    # Reductions section load 
    reductions_data = group[['Field Name (MRV)', 'reduced']].copy()
    reductions_data.rename(columns={'Field Name (MRV)': 'Field'}, inplace=True)
    total_reductions = reductions_data['reduced'].sum()
    # Removals section load 
    removals_data = group[['Field Name (MRV)', 'baseline_dsoc', 'dsoc', 'removed']].copy()
    removals_data.rename(columns={'Field Name (MRV)': 'Field', 'removed': 'Removed'}, inplace=True)
    total_removals = removals_data['Removed'].sum()

    # bulleted results 
    #doc.add_paragraph(f"Total Reductions: {total_reductions:.3f} tonnes CO2e", style='List Bullet')
    #doc.add_paragraph(f"Total Removals: {total_removals:.3f} tonnes CO2e", style='List Bullet')
    #doc.add_paragraph(f"Total Acres: {fields_data['Acres'].sum():.3f}", style='List Bullet')


    # Create a bullet point for Total Reductions and apply styles
    add_custom_bullet(doc, "Total Reductions: ", f"{total_reductions:.3f} tonnes CO2e")

    # Create a bullet point for Total Removals and apply styles
    add_custom_bullet(doc, "Total Removals: ", f"{total_removals:.3f} tonnes CO2e")

    # Create a bullet point for Total Acres and apply styles
    add_custom_bullet(doc, "Total Acres: ", f"{fields_data['Acres'].sum():.3f}")



    # Fields section
    create_table(doc, fields_data.values, ['Field', 'Acres', 'Commodity'], title='Fields')
    doc.add_paragraph()
    para = doc.add_paragraph()
    run = para.add_run(f"Total Acres: {fields_data['Acres'].sum():.2f}")
    run.italic = True

    # Practices section
    practices_data = group[['Field Name (MRV)', 'Practice Change']].copy()
    practices_data.rename(columns={'Field Name (MRV)': 'Field'}, inplace=True)
    create_table(doc, practices_data.values, ['Field', 'Practice Change'], title='Practices')

    # Soil Data section
    soil_data = group[['Field Name (MRV)', 'soil_avg_soc', 'soil_avg_ph', 'soil_avg_bulkdensity', 'soil_clay_fraction']].copy()
    soil_data.rename(columns={'Field Name (MRV)': 'Field', 'soil_avg_soc': 'SOC', 'soil_avg_ph': 'pH', 'soil_avg_bulkdensity': 'BD', 'soil_clay_fraction': 'Clay'}, inplace=True)
    create_table(doc, soil_data.values, ['Field', 'SOC', 'pH', 'BD', 'Clay'], title='Soil Data')

    # Reductions section
    doc.add_heading('Reductions', level=1)
    doc.add_paragraph("This section shows the GHG emissions reductions from the intervention, and includes DNDC modeled reductions as well as reductions from field activity changes (calculated through ecoinvent3.8 emission factors).")
    
    # Example for Total Reductions
    create_table(doc, reductions_data.values, ['Field', 'Reduced'], title='Total Reductions')
    doc.add_paragraph()
    para = doc.add_paragraph()
    run = para.add_run(f"Total Reductions: {total_reductions:.3f} tonnes CO2e")
    run.italic = True

    # Similar approach for N2O, Methane, etc.
    
    # Direct N2O
    group['Delta Direct N2O'] = calculate_delta(group, 'baseline_n20_direct', 'n2o_direct')
    # N2O Direct section
    n2o_data = group[['Field Name (MRV)', 'baseline_n20_direct', 'n2o_direct', 'Delta Direct N2O']].copy()
    n2o_data.rename(columns={
        'Field Name (MRV)': 'Field',
        'baseline_n20_direct': 'Direct N2O Baseline',
        'n2o_direct': 'Direct N2O Practice',
        'Delta Direct N2O': 'Delta'
    }, inplace=True)
    create_table(doc, n2o_data.values, ['Field', 'Direct N2O Baseline', 'Direct N2O Practice', 'Delta'], title='N2O Direct Emissions')
    doc.add_paragraph()
    # Summarize the total Delta Direct N2O
    total_delta_direct_n2o = n2o_data['Delta'].sum()
    para = doc.add_paragraph()
    run = para.add_run(f"Total Direct N2O Reductions: {total_delta_direct_n2o:.3f} tonnes CO2e")
    run.italic = True

    #Indirect N2O
    group['Delta Indirect N2O'] = calculate_delta(group, 'baseline_n2o_indirect', 'n2o_indirect')
    # N2O Direct section
    n2oindirect_data = group[['Field Name (MRV)', 'baseline_n2o_indirect', 'n2o_indirect', 'Delta Indirect N2O']].copy()
    n2oindirect_data.rename(columns={
        'Field Name (MRV)': 'Field',
        'baseline_n2o_indirect': 'Indirect N2O Baseline',
        'n2o_indirect': 'Indirect N2O Practice',
        'Delta Indirect N2O': 'Delta'
    }, inplace=True)
    create_table(doc, n2oindirect_data.values, ['Field', 'Indirect N2O Baseline', 'Indirect N2O Practice', 'Delta'], title='N2O Indirect Emissions')
    doc.add_paragraph()
    # Summarize the total Delta Indirect N2O
    total_delta_indirect_n2o = n2oindirect_data['Delta'].sum()
    para = doc.add_paragraph()
    run = para.add_run(f"Total Indirect N2O Reductions: {total_delta_indirect_n2o:.3f} tonnes CO2e")
    run.italic = True

    # Methane 
    group['Delta Methane'] = calculate_delta(group, 'baseline_methane', 'methane')
    # Methane section
    methane_data = group[['Field Name (MRV)', 'baseline_methane', 'methane', 'Delta Methane']].copy()
    methane_data.rename(columns={
        'Field Name (MRV)': 'Field',
        'baseline_methane': 'CH4 Baseline',
        'methane': 'CH4 Practice',
        'Delta Methane': 'Delta'
    }, inplace=True)
    create_table(doc, methane_data.values, ['Field', 'CH4 Baseline', 'CH4 Practice', 'Delta'], title='Methane Emissions')
    doc.add_paragraph()
    # Summarize the total Methane
    total_delta_methane = methane_data['Delta'].sum()
    para = doc.add_paragraph()
    run = para.add_run(f"Total Methane Reductions: {total_delta_methane:.3f} tonnes CO2e")
    run.italic = True

    #Emission Factors 
    group['Delta Field Emissions'] = calculate_delta(group, 'field_baseline_emissions_demands_fossil', 'field_practice_emissions_demands_fossil')
    # Emissions section
    emissions_data = group[['Field Name (MRV)', 'field_baseline_emissions_demands_fossil', 'field_practice_emissions_demands_fossil', 'Delta Field Emissions']].copy()
    emissions_data.rename(columns={
        'Field Name (MRV)': 'Field',
        'field_baseline_emissions_demands_fossil': 'Emissions Demands Baseline',
        'field_practice_emissions_demands_fossil': 'Emissions Demands Practice',
        'Delta Field Emissions': 'Delta'
    }, inplace=True)
    create_table(doc, emissions_data.values, ['Field', 'Emissions Demands Baseline', 'Emissions Demands Practice', 'Delta'], title='Emissions Demands')
    doc.add_paragraph()
    # Summarize the total Emissions
    total_delta_emissions = emissions_data['Delta'].sum()
    para = doc.add_paragraph()
    doc.add_paragraph("Derived from ecoinvent3.8 cutoff emission factors for upstream and on-field process emissions.")
    run = para.add_run(f"Total Emissions Demands Reductions: {total_delta_emissions:.3f} kg CO2e")
    run.italic = True

    # Removals section
    doc.add_heading('Removals', level=1)
    doc.add_paragraph("This section shows the carbon removal from the intervention as modeled by the DNDC.")
    create_table(doc, removals_data.values, ['Field', 'Baseline DSOC', 'DSOC', 'Removed'], title='Total Removals')
    doc.add_paragraph()
    para = doc.add_paragraph()
    run = para.add_run(f"Total Removals: {total_removals:.3f} tonnes CO2e")
    run.italic = True

    # Save the document
    safe_producer_name = producer.replace('/', '_').replace('\\', '_')
    doc_file_name = f"{safe_producer_name}.docx"
    doc.save(os.path.join(output_dir, doc_file_name))

print("Reports generated successfully.")

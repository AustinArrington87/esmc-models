import pandas as pd
from docx import Document
from docx.shared import Inches, Pt
from docx.oxml import OxmlElement
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.text import WD_UNDERLINE
from docx.shared import RGBColor
from docx.oxml.ns import qn
import os

# Year generating reports for
run_year = "2022"

# Load the producer mapping
mapping_df = pd.read_csv('Mapping.csv')
producer_mapping = dict(zip(mapping_df['project_producer'], mapping_df['display_name']))

def load_data(year, batch):
    file_path = f"{year}_{batch}.csv" if batch == "Quantified" else f"{year}_{batch}.csv"
    return pd.read_csv(file_path)

def format_cell_data(data):
    try:
        float_data = float(data)
        return f"{float_data:.3f}"
    except ValueError:
        return data

def create_table(doc, data, headers, font_name, font_size, font_color, scale_emissions=False):
    table = doc.add_table(rows=len(data) + 1, cols=len(headers))
    table.style = 'Table Grid'

    for i, header in enumerate(headers):
        cell = table.cell(0, i)
        run = cell.paragraphs[0].add_run(header)
        run.font.name = font_name
        run.font.size = font_size
        run.font.color.rgb = font_color
        run.font.bold = True
        cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    for row_idx, row_data in enumerate(data, start=1):
        for col_idx, cell_data in enumerate(row_data):
            if scale_emissions and isinstance(cell_data, (int, float)):
                cell_data *= 0.001
            if isinstance(cell_data, float):
                cell_data = f"{cell_data:.3f}"
            cell = table.cell(row_idx, col_idx)
            run = cell.paragraphs[0].add_run(str(cell_data))
            run.font.name = font_name
            run.font.size = font_size
            run.font.color.rgb = font_color
            cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

def calculate_delta(group, baseline_column, practice_column):
    return group[baseline_column] - group[practice_column]

def add_custom_bullet(doc, label, value):
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(label)
    run.bold = True
    run.font.name = 'Calibri'
    run.font.size = Pt(12)

    value_run = p.add_run(value)
    value_run.font.name = 'Calibri'
    value_run.font.size = Pt(12)

def add_items_bulleted_list(document, practices):
    for practice in practices:
        paragraph = document.add_paragraph(style='List Bullet')
        run = paragraph.add_run(practice)
        run.font.name = 'Calibri'
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(0, 0, 0)
        # Indent the paragraph
        paragraph_format = paragraph.paragraph_format
        paragraph_format.left_indent = Inches(0.5)


def add_hyperlink(paragraph, text, url):
    part = paragraph.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)

    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)
    hyperlink.set(qn('w:history'), '1')

    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    c = OxmlElement('w:color')
    c.set(qn('w:val'), '0000FF')
    rPr.append(c)
    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'single')
    rPr.append(u)
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)

    return hyperlink

def initialize_document():
    doc = Document()
    # Add header image only once
    section = doc.sections[0]
    header = section.header
    header_para = header.paragraphs[0]
    run = header_para.add_run()
    run.add_picture('header.png', width=Inches(6))
    return doc

def generate_project_report(data, batch, doc):
    font_name = 'Calibri'
    font_size = Pt(12)
    font_color = RGBColor(0, 0, 0)

    for project, group in data.groupby('Project'):
        if project not in doc.projects:
            doc.projects[project] = initialize_document()

        document = doc.projects[project]

        document.add_paragraph()  # empty space
        para = document.add_heading()
        run = para.add_run(f"Eco-Harvest Carbon Impact Report")
        run_font = run.font
        run_font.name = 'Calibri'
        run_font.size = Pt(24)
        run_font.underline = WD_UNDERLINE.SINGLE
        run_font.color.rgb = RGBColor(0, 0, 0)

        # summary info at top
        summary_para1 = document.add_paragraph()
        run_bold = summary_para1.add_run('Project: ')
        run_bold.font.bold = True
        run_bold.font.name = 'Calibri'
        run_bold.font.size = Pt(12)
        run_bold.font.color.rgb = RGBColor(0, 0, 0)
        run_regular = summary_para1.add_run(project)
        run_regular.font.name = 'Calibri'
        run_regular.font.size = Pt(12)
        run_regular.font.color.rgb = RGBColor(0, 0, 0)

        summary_para2 = document.add_paragraph()
        run_bold = summary_para2.add_run('Year: ')
        run_bold.font.bold = True
        run_bold.font.name = 'Calibri'
        run_bold.font.size = Pt(12)
        run_bold.font.color.rgb = RGBColor(0, 0, 0)
        run_regular = summary_para2.add_run(run_year)
        run_regular.font.name = 'Calibri'
        run_regular.font.size = Pt(12)
        run_regular.font.color.rgb = RGBColor(0, 0, 0)

        #Batch
        summary_para3 = document.add_paragraph()
        run_bold = summary_para3.add_run('Batch: ')
        run_bold.font.bold = True
        run_bold.font.name = 'Calibri'
        run_bold.font.size = Pt(12)
        run_bold.font.color.rgb = RGBColor(0, 0, 0)
        run_regular = summary_para3.add_run(batch)
        run_regular.font.name = 'Calibri'
        run_regular.font.size = Pt(12)
        run_regular.font.color.rgb = RGBColor(0, 0, 0)

        # load data 
        fields_data = group[['Field Name (MRV)', 'Acres', 'Commodity', 'Producer (Project)', 'Practice Change']].copy()
        fields_data.rename(columns={'Field Name (MRV)': 'Field'}, inplace=True)

        # Calculate totals
        total_acres = fields_data['Acres'].sum()
        total_producers = fields_data['Producer (Project)'].nunique()
        total_fields = fields_data['Field'].nunique()
        total_crops = list(set(fields_data['Commodity']))
        total_crops = [x for x in total_crops if str(x) != 'nan']
        total_practices = list(set(fields_data['Practice Change']))
        total_practices = [x for x in total_practices if str(x) != 'nan']

        # Reductions data
        reductions_data = group[['Field Name (MRV)', 'Practice Change', 'reduced' if batch == "Quantified" else 'reduced_adjusted']].copy()
        reductions_data.rename(columns={'Field Name (MRV)': 'Field', 'Practice Change': 'Practice_Change'}, inplace=True)
        reductions_data[reductions_data.columns[2]] = pd.to_numeric(reductions_data[reductions_data.columns[2]], errors='coerce')
        total_reductions = reductions_data[reductions_data.columns[2]].sum()

        # Removals data
        removals_data = group[['Field Name (MRV)', 'Practice Change', 'baseline_dsoc', 'dsoc', 'removed']].copy()
        removals_data.rename(columns={'Field Name (MRV)': 'Field', 'Practice Change': 'Practice_Change', 'removed': 'Removed'}, inplace=True)
        removals_data['Removed'] = pd.to_numeric(removals_data['Removed'], errors='coerce')
        total_removals = removals_data['Removed'].sum()
        total_carbon_g = float(total_reductions) + float(total_removals)

        # Calculate mean and standard deviation
        mean_reductions = reductions_data[reductions_data.columns[2]].mean()
        std_reductions = reductions_data[reductions_data.columns[2]].std()

        mean_removals = removals_data['Removed'].mean()
        std_removals = removals_data['Removed'].std()

        # Aggregate outcomes by practice change
        practice_stats = []
        for practice in total_practices:
            practice_reductions = reductions_data[reductions_data['Practice_Change'] == practice]
            practice_removals = removals_data[removals_data['Practice_Change'] == practice]
            practice_acres = fields_data[fields_data['Practice Change'] == practice]['Acres'].sum()

            mean_reductions = practice_reductions[practice_reductions.columns[2]].mean()
            std_reductions = practice_reductions[practice_reductions.columns[2]].std()

            mean_removals = practice_removals['Removed'].mean()
            std_removals = practice_removals['Removed'].std()

            practice_stats.append((practice, mean_reductions, std_reductions, mean_removals, std_removals, practice_acres))


        #Producers
        summary_para4 = document.add_paragraph()
        run_bold = summary_para4.add_run('Producers: ')
        run_bold.font.bold = True
        run_bold.font.name = 'Calibri'
        run_bold.font.size = Pt(12)
        run_bold.font.color.rgb = RGBColor(0, 0, 0)
        run_regular = summary_para4.add_run(str(total_producers))
        run_regular.font.name = 'Calibri'
        run_regular.font.size = Pt(12)
        run_regular.font.color.rgb = RGBColor(0, 0, 0)

        # Acres 
        summary_para5 = document.add_paragraph()
        run_bold = summary_para5.add_run('Acres: ')
        run_bold.font.bold = True
        run_bold.font.name = 'Calibri'
        run_bold.font.size = Pt(12)
        run_bold.font.color.rgb = RGBColor(0, 0, 0)
        run_regular = summary_para5.add_run(str(round(total_acres,3)))
        run_regular.font.name = 'Calibri'
        run_regular.font.size = Pt(12)
        run_regular.font.color.rgb = RGBColor(0, 0, 0)

        intro_para = document.add_paragraph()
        run_intro = intro_para.add_run("Practice Change(s):")
        run_intro.font.bold = True
        run_intro.font.name = 'Calibri'
        run_intro.font.size = Pt(12)
        run_intro.font.color.rgb = RGBColor(0, 0, 0)
        add_items_bulleted_list(document, total_practices)

        intro_para = document.add_paragraph()
        run_intro = intro_para.add_run("Crop(s):")
        run_intro.font.bold = True
        run_intro.font.name = 'Calibri'
        run_intro.font.size = Pt(12)
        run_intro.font.color.rgb = RGBColor(0, 0, 0)
        add_items_bulleted_list(document, total_crops)


        para = document.add_heading()
        run = para.add_run(f"Eco-Harvest Results Overview")
        run_font = run.font
        run_font.name = 'Calibri'
        run_font.size = Pt(16)
        run_font.underline = WD_UNDERLINE.SINGLE
        run_font.color.rgb = RGBColor(0, 0, 0)

        # Payment logic
        specific_projects_16 = [
            "AGI SGP Market", "AGI SGP Pipeline", "NACD SGP Market",
            "Northern Plains Cropping", "NGP-NACD Market", "NGP-NACD Pipeline"
        ]
        specific_projects_20 = ["TNC Minnesota", "TNC Nebraska"]

        payment_amount = 0
        payment_message = "Producers in this project received a payment of $X/acre OR $X/impact."

        if project in specific_projects_16:
            payment_amount = 16 * total_carbon_g
            if payment_amount < 0:
                payment_message = "These fields did not accrue enough impact for payment this growing season."
            else:
                payment_message = f"Producers in this project received a payment of $16/outcome. Total payment is ${payment_amount:.2f}."
        elif project in specific_projects_20:
            payment_amount = 20 * total_acres
            if payment_amount < 0:
                payment_message = "These fields did not accrue enough impact for payment this growing season."
            else:
                payment_message = f"Producers in this project received a payment of $20/acre. Total payment is ${payment_amount:.2f}."


        add_custom_bullet(document, "Total Carbon Impacts (Removals + Reductions) Generated: ", f"{total_carbon_g:.3f} metric tonnes of carbon dioxide equivalent (mtCO2e)")
        add_custom_bullet(document, "Emission Reductions: ", f"Practice changes in the project avoided {total_reductions:.3f} metric tonnes of carbon dioxide equivalent (mtCO2e) reductions.")
        add_custom_bullet(document, "Carbon Removals: ", f"Soils removed {total_removals:.3f} mtCO2e due to producer's new practice changes.")
        add_custom_bullet(document, "Total Payment for Outcomes/Impact Units: ", f"{payment_message}")


        summary_para4 = document.add_paragraph()
        run_regular = summary_para4.add_run("*ESMC requires complete historical data to generate a baseline, so fields with historical data gaps are not modeled.")
        run_regular.font.name = 'Calibri'
        run_regular.font.size = Pt(9)
        run_regular.font.color.rgb = RGBColor(0, 0, 0)

        # REWARDS
        para = document.add_heading()
        run = para.add_run(f"Eco-Harvest: rewarding producers for regenerative ag")
        run_font = run.font
        run_font.name = 'Calibri'
        run_font.size = Pt(16)
        run_font.underline = WD_UNDERLINE.SINGLE
        run_font.color.rgb = RGBColor(0, 0, 0)

        add_custom_bullet(document, "Customizable practices: ", f"Choose how to implement regenerative practices.")
        add_custom_bullet(document, "Fair Payments: ", f"As a non-profit, ESMC maximizes producer benefits.")
        add_custom_bullet(document, "Return on Investment: ", f"Many farmers experience a 15-25% return on investment after 3-5 years of regenerative farming.")
        add_custom_bullet(document, "Weather-Ready Farming: ", f"Increased weather resilience, soil quality, and efficiency")

        thankyou_para = document.add_paragraph()
        run_regular = thankyou_para.add_run("Thank you for participating in the Eco-Harvest program and for your dedication to regenerative agriculture.")
        run_regular.font.name = 'Calibri'
        run_regular.font.size = Pt(12)
        run_regular.font.color.rgb = RGBColor(0, 0, 0)

        # Appendix
        para = document.add_heading()
        run = para.add_run(f"Appendix")
        run_font = run.font
        run_font.name = 'Calibri'
        run_font.size = Pt(16)
        run_font.underline = WD_UNDERLINE.SINGLE
        run_font.color.rgb = RGBColor(0, 0, 0)

        
        para = document.add_heading()
        run = para.add_run(f"Carbon Impact Explained")
        run_font = run.font
        run_font.name = 'Calibri'
        run_font.size = Pt(14)
        run_font.underline = WD_UNDERLINE.SINGLE
        run_font.color.rgb = RGBColor(0, 0, 0)
        summary_para7 = document.add_paragraph()
        run_regular = summary_para7.add_run("ESMC uses historical data, a scientific model, soil samples and weather data to calculate emissions reductions and soil carbon removals. Any negative emissions are increased emissions, which is not unusual in the early years of adopting a practice change. If concerned, please speak with your technical advisor about further modifications to your practice change(s).")
        run_regular.font.name = 'Calibri'
        run_regular.font.size = Pt(12)
        run_regular.font.color.rgb = RGBColor(0, 0, 0)

        add_custom_bullet(document, "Greenhouse Gas Emissions (GHG) Reductions (mtCO2e) ", f"are the sum of the reductions of methane (CH4), direct nitrous oxide (N2O), indirect nitrous oxide (N2O) and supply chain and operational emissions** between the baseline year and the project year.")
        add_custom_bullet(document, "Greenhouse Gas Removals ", f"is the difference in the amount of carbon in the soil between the baseline year and the project year.")
        add_custom_bullet(document, "Verified Carbon Impacts ", f"are impacts which were verified by our third-party auditor, SustainCert")
        add_custom_bullet(document, "Quantified Carbon Impacts ", f"are impacts which were verified by ESMC QA/QC, satellite data, and had most but not all data points supplied by the producer. Since some data points were not from the producer, the quantified carbon impacts are not eligible to be verified.")

        summary_para8 = document.add_paragraph()
        run_regular = summary_para8.add_run("**Upstream and on-field process emissions associated with management such as fuel used on farm and imbedded emissions of fertilizer products; ESMC utilized emission factors from Ecoinvent and WFLD.")
        run_regular.font.name = 'Calibri'
        run_regular.font.size = Pt(9)
        run_regular.font.color.rgb = RGBColor(0, 0, 0)

        group['Delta Direct N2O'] = calculate_delta(group, 'baseline_n20_direct', 'n2o_direct')
        n2o_data = group[['Field Name (MRV)', 'baseline_n20_direct', 'n2o_direct', 'Delta Direct N2O']].copy()
        n2o_data.rename(columns={
            'Field Name (MRV)': 'Field',
            'baseline_n20_direct': 'Direct N2O Baseline',
            'n2o_direct': 'Direct N2O Practice',
            'Delta Direct N2O': 'Delta'
        }, inplace=True)
        total_delta_direct_n2o = n2o_data['Delta'].sum()

        group['Delta Indirect N2O'] = calculate_delta(group, 'baseline_n2o_indirect' if batch == "Quantified" else 'baseline_n2o_indirect_adjusted', 'n2o_indirect' if batch == "Quantified" else 'n2o_indirect_adjusted')
        n2oindirect_data = group[['Field Name (MRV)', 'baseline_n2o_indirect', 'n2o_indirect', 'Delta Indirect N2O']].copy()
        n2oindirect_data.rename(columns={
            'Field Name (MRV)': 'Field',
            'baseline_n2o_indirect' if batch == "Quantified" else 'baseline_n2o_indirect_adjusted': 'Indirect N2O Baseline',
            'n2o_indirect' if batch == "Quantified" else 'n2o_indirect_adjusted': 'Indirect N2O Practice',
            'Delta Indirect N2O': 'Delta'
        }, inplace=True)
        total_delta_indirect_n2o = n2oindirect_data['Delta'].sum()

        group['Delta Methane'] = calculate_delta(group, 'baseline_methane' if batch == "Quantified" else 'baseline_methane_adjusted', 'methane' if batch == "Quantified" else 'methane_adjusted')
        methane_data = group[['Field Name (MRV)', 'baseline_methane', 'methane', 'Delta Methane']].copy()
        methane_data.rename(columns={
            'Field Name (MRV)': 'Field',
            'baseline_methane': 'CH4 Baseline',
            'methane': 'CH4 Practice',
            'Delta Methane': 'Delta'
        }, inplace=True)
        total_delta_methane = methane_data['Delta'].sum()

        group['Delta Field Emissions'] = calculate_delta(group, 'field_baseline_emissions', 'field_practice_emissions')
        emissions_data = group[['Field Name (MRV)', 'field_baseline_emissions', 'field_practice_emissions', 'Delta Field Emissions']].copy()
        emissions_data.rename(columns={
            'Field Name (MRV)': 'Field',
            'field_baseline_emissions': 'Supply Chain & Operational Emissions Baseline',
            'field_practice_emissions': 'Supply Chain & Operational Emissions Practice',
            'Delta Field Emissions': 'Delta'
        }, inplace=True)
        total_delta_emissions = emissions_data['Delta'].sum() * 0.001

        impact_table_data = [
            ("Greenhouse Gas Emissions Reductions", total_reductions / total_acres if total_acres else 0, total_reductions),
            ("N2O Indirect", total_delta_indirect_n2o / total_acres if total_acres else 0, total_delta_indirect_n2o),
            ("N2O Direct", total_delta_direct_n2o / total_acres if total_acres else 0, total_delta_direct_n2o),
            ("Methane", total_delta_methane / total_acres if total_acres else 0, total_delta_methane),
            ("Supply Chain and Operational Emissions", total_delta_emissions / total_acres if total_acres else 0, total_delta_emissions),
            ("Total Soil Carbon Removals", total_removals / total_acres if total_acres else 0, total_removals),
            ("Total Carbon Impact (Reductions + Removals)", total_carbon_g / total_acres if total_acres else 0, total_carbon_g)
        ]


         # PROJECT LEVEL ANALYSIS
        para = document.add_heading()
        run = para.add_run(f"Project Level Analysis")
        run_font = run.font
        run_font.name = 'Calibri'
        run_font.size = Pt(14)
        run_font.underline = WD_UNDERLINE.SINGLE
        run_font.color.rgb = RGBColor(0, 0, 0)

        # Summary table 

        create_table(document, impact_table_data, ['Impact Breakdown', 'mtCO2e (delta) per acre', 'mtCO2e (delta) total'], font_name, font_size, font_color)

        # Stats 

        summary_space = document.add_paragraph()
        summary_para9 = document.add_paragraph()
        run_regular = summary_para9.add_run("Mean and Standard Deviation across project as a whole.")
        run_regular.font.name = 'Calibri'
        run_regular.font.size = Pt(12)
        run_regular.font.color.rgb = RGBColor(0, 0, 0)

        summary_para10 = document.add_paragraph()
        run_regular = summary_para10.add_run("For removals and reductions:")
        run_regular.font.name = 'Calibri'
        run_regular.font.size = Pt(12)
        run_regular.font.color.rgb = RGBColor(0, 0, 0)
        run_regular.italic = True

        # Prepare statistics table data
        stats_table_data = [
            ("Removals", mean_removals, std_removals),
            ("Reductions", mean_reductions, std_reductions)
        ]

        # Create table with Mean and Standard Deviation
        create_table(document, stats_table_data, ['Outcome', 'Mean (mtCO2e)', 'Standard Deviation (mtCO2e)'], font_name, font_size, font_color)

        summary_para10 = document.add_paragraph()
        run_regular = summary_para10.add_run("By practice change(s) implemented:")
        run_regular.font.name = 'Calibri'
        run_regular.font.size = Pt(12)
        run_regular.font.color.rgb = RGBColor(0, 0, 0)
        run_regular.italic = True

        # Prepare practice change stats table data
        stats_table_data = [
            (practice, mean_reductions, std_reductions, mean_removals, std_removals, practice_acres)
            for practice, mean_reductions, std_reductions, mean_removals, std_removals, practice_acres in practice_stats
        ]

        # Create table with Mean and Standard Deviation segmented by practice change
        create_table(document, stats_table_data, ['Practice Change', 'Mean Reductions (mtCO2e)', 'Std Dev Reductions (mtCO2e)', 'Mean Removals (mtCO2e)', 'Std Dev Removals (mtCO2e)', 'Acres'], font_name, font_size, font_color)

        # Table Segmented by Crops

        # Reductions data
        reductions_data = group[['Field Name (MRV)', 'Commodity', 'reduced' if batch == "Quantified" else 'reduced_adjusted']].copy()
        reductions_data.rename(columns={'Field Name (MRV)': 'Field'}, inplace=True)
        reductions_data[reductions_data.columns[2]] = pd.to_numeric(reductions_data[reductions_data.columns[2]], errors='coerce')
        total_reductions = reductions_data[reductions_data.columns[2]].sum()

        # Removals data
        removals_data = group[['Field Name (MRV)', 'Commodity', 'baseline_dsoc', 'dsoc', 'removed']].copy()
        removals_data.rename(columns={'Field Name (MRV)': 'Field', 'removed': 'Removed'}, inplace=True)
        removals_data['Removed'] = pd.to_numeric(removals_data['Removed'], errors='coerce')
        total_removals = removals_data['Removed'].sum()
        total_carbon_g = float(total_reductions) + float(total_removals)

        # Ensure commodity column is treated as string
        reductions_data['Commodity'] = reductions_data['Commodity'].astype(str)
        removals_data['Commodity'] = removals_data['Commodity'].astype(str)

        # Aggregate outcomes by practice change
        commodity_stats = []
        for commodity in total_crops:
            commodity_reductions = reductions_data[reductions_data['Commodity'] == commodity]
            commodity_removals = removals_data[removals_data['Commodity'] == commodity]
            commodity_acres = fields_data[fields_data['Commodity'] == commodity]['Acres'].sum()

            mean_reductions = commodity_reductions[commodity_reductions.columns[2]].mean()
            std_reductions = commodity_reductions[commodity_reductions.columns[2]].std()

            mean_removals = commodity_removals['Removed'].mean()
            std_removals = commodity_removals['Removed'].std()

            commodity_stats.append((commodity, mean_reductions, std_reductions, mean_removals, std_removals, commodity_acres))

        summary_para10 = document.add_paragraph()
        run_regular = summary_para10.add_run("By crop(s) planted:")
        run_regular.font.name = 'Calibri'
        run_regular.font.size = Pt(12)
        run_regular.font.color.rgb = RGBColor(0, 0, 0)
        run_regular.italic = True

        # Prepare practice change stats table data
        stats_table_data = [
            (commodity, mean_reductions, std_reductions, mean_removals, std_removals, commodity_acres)
            for commodity, mean_reductions, std_reductions, mean_removals, std_removals, commodity_acres in commodity_stats
        ]

        # Create table with Mean and Standard Deviation segmented by practice change
        create_table(document, stats_table_data, ['Commodity Crop', 'Mean Reductions (mtCO2e)', 'Std Dev Reductions (mtCO2e)', 'Mean Removals (mtCO2e)', 'Std Dev Removals (mtCO2e)', 'Acres'], font_name, font_size, font_color)

        ######## 


        document.add_picture('diagram.png')
        italic_paragraph = document.add_paragraph()
        italic_run = italic_paragraph.add_run('Figure 1: Total impact is the combined value of GHG emissions reductions and soil carbon stored. This diagram illustrates a general improvement from the baseline to project year – GHG emissions decrease and soil carbon increases.')
        italic_run.italic = True
        italic_run.font_name = 'Calibri'
        italic_run.font.size = Pt(11)

        document.add_paragraph()
        para = document.add_paragraph()
        run = para.add_run("Real-world Comparisons")
        run.bold = True
        run.font.name = 'Calibri'

        summary_para9 = document.add_paragraph()
        run_regular = summary_para9.add_run("Use ")
        run_regular.font.name = 'Calibri'
        run_regular.font.size = Pt(12)
        run_regular.font.color.rgb = RGBColor(0, 0, 0)
        add_hyperlink(summary_para9, 'this tool', 'https://www.epa.gov/energy/greenhouse-gas-equivalencies-calculator#resultsg')
        run_regular = summary_para9.add_run(" to compare carbon impacts to everyday examples.")
        run_regular.font.name = 'Calibri'
        run_regular.font.size = Pt(12)
        run_regular.font.color.rgb = RGBColor(0, 0, 0)

    return doc.projects

# Load data
verified_data = load_data(run_year, "Verified")
quantified_data = load_data(run_year, "Quantified")

# Initialize document dictionary
class DocumentHolder:
    def __init__(self):
        self.projects = {}

doc = DocumentHolder()

# Generate reports for each batch
generate_project_report(verified_data, "Verified", doc)
generate_project_report(quantified_data, "Quantified", doc)

# Save documents
output_dir = "Projects"
if not os.path.exists(output_dir):
    os.makedirs(output_dir)

for project, document in doc.projects.items():
    safe_project_name = project.replace('/', '_').replace('\\', '_')
    doc_file_name = f"{safe_project_name}.docx"
    document.save(os.path.join(output_dir, doc_file_name))

print("Project reports generated successfully.")

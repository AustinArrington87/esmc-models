import pandas as pd
from docx import Document
from docx.shared import Inches, Pt
from docx.oxml import OxmlElement
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.text import WD_UNDERLINE
from docx.shared import RGBColor
from docx.oxml.ns import qn
import os

# year genearating reports for
run_year = "2023"
# used for mapping anonymized prooducer name 
mapping_df = pd.read_csv('Mapping.csv')
producer_mapping = dict(zip(mapping_df['project_producer'], mapping_df['display_name']))

def load_data(year, batch):
    file_path = f"{year}_{batch}.csv" if batch == "Quantified" else f"{year}_{batch}.csv"
    return pd.read_csv(file_path)

def format_cell_data(data):
    try:
        float_data = float(data)
        return f"{float_data:.3f}"
    except ValueError:
        return data

def create_table(doc, data, headers, font_name, font_size, font_color, scale_emissions=False):
    table = doc.add_table(rows=len(data) + 1, cols=len(headers))
    table.style = 'Table Grid'

    for i, header in enumerate(headers):
        cell = table.cell(0, i)
        run = cell.paragraphs[0].add_run(header)
        run.font.name = font_name
        run.font.size = font_size
        run.font.color.rgb = font_color
        run.font.bold = True
        cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    for row_idx, row_data in enumerate(data, start=1):
        for col_idx, cell_data in enumerate(row_data):
            if scale_emissions and isinstance(cell_data, (int, float)):
                cell_data *= 0.001
            if isinstance(cell_data, float):
                cell_data = f"{cell_data:.3f}"
            cell = table.cell(row_idx, col_idx)
            run = cell.paragraphs[0].add_run(str(cell_data))
            run.font.name = font_name
            run.font.size = font_size
            run.font.color.rgb = font_color
            cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

def calculate_delta(group, baseline_column, practice_column):
    return group[baseline_column] - group[practice_column]

def add_custom_bullet(doc, label, value):
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(label)
    run.bold = True
    run.font.name = 'Calibri'
    run.font.size = Pt(12)

    value_run = p.add_run(value)
    value_run.font.name = 'Calibri'
    value_run.font.size = Pt(12)

def add_hyperlink(paragraph, text, url):
    part = paragraph.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)

    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)
    hyperlink.set(qn('w:history'), '1')

    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    c = OxmlElement('w:color')
    c.set(qn('w:val'), '0000FF')
    rPr.append(c)
    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'single')
    rPr.append(u)
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)

    return hyperlink

def initialize_document():
    doc = Document()
    # Add header image only once
    section = doc.sections[0]
    header = section.header
    header_para = header.paragraphs[0]
    run = header_para.add_run()
    run.add_picture('header.png', width=Inches(6))
    return doc

def generate_report(data, batch, doc):
    font_name = 'Calibri'
    font_size = Pt(12)
    font_color = RGBColor(0, 0, 0)

    for producer, group in data.groupby('Producer (Project)'):
        if producer not in doc.producers:
            doc.producers[producer] = initialize_document()

        document = doc.producers[producer]

        project = group['Project'].iloc[0]

        document.add_paragraph()  # empty space
        para = document.add_heading()
        run = para.add_run(f"Eco-Harvest Carbon Impact Report")
        run_font = run.font
        run_font.name = 'Calibri'
        run_font.size = Pt(24)
        run_font.underline = WD_UNDERLINE.SINGLE
        run_font.color.rgb = RGBColor(0, 0, 0)

        # summary info at top
        summary_para1 = document.add_paragraph()
        run_bold = summary_para1.add_run('Producer: ')
        run_bold.font.bold = True
        run_bold.font.name = 'Calibri'
        run_bold.font.size = Pt(12)
        run_bold.font.color.rgb = RGBColor(0, 0, 0)
        try:
            display_name = producer_mapping.get(producer, producer)
            run_regular = summary_para1.add_run(f"{display_name} ({producer})")
        except:
            run_regular = summary_para1.add_run(producer)
        run_regular.font.name = 'Calibri'
        run_regular.font.size = Pt(12)
        run_regular.font.color.rgb = RGBColor(0, 0, 0)

        summary_para2 = document.add_paragraph()
        run_bold = summary_para2.add_run('Project: ')
        run_bold.font.bold = True
        run_bold.font.name = 'Calibri'
        run_bold.font.size = Pt(12)
        run_bold.font.color.rgb = RGBColor(0, 0, 0)
        run_regular = summary_para2.add_run(project)
        run_regular.font.name = 'Calibri'
        run_regular.font.size = Pt(12)
        run_regular.font.color.rgb = RGBColor(0, 0, 0)

        summary_para3 = document.add_paragraph()
        run_bold = summary_para3.add_run('Year: ')
        run_bold.font.bold = True
        run_bold.font.name = 'Calibri'
        run_bold.font.size = Pt(12)
        run_bold.font.color.rgb = RGBColor(0, 0, 0)
        run_regular = summary_para3.add_run(run_year)
        run_regular.font.name = 'Calibri'
        run_regular.font.size = Pt(12)
        run_regular.font.color.rgb = RGBColor(0, 0, 0)

        summary_para4 = document.add_paragraph()
        run_bold = summary_para4.add_run('Batch: ')
        run_bold.font.bold = True
        run_bold.font.name = 'Calibri'
        run_bold.font.size = Pt(12)
        run_bold.font.color.rgb = RGBColor(0, 0, 0)
        run_regular = summary_para4.add_run(batch)
        run_regular.font.name = 'Calibri'
        run_regular.font.size = Pt(12)
        run_regular.font.color.rgb = RGBColor(0, 0, 0)

        para = document.add_heading()
        run = para.add_run(f"Your Results")
        run_font = run.font
        run_font.name = 'Calibri'
        run_font.size = Pt(16)
        run_font.underline = WD_UNDERLINE.SINGLE
        run_font.color.rgb = RGBColor(0, 0, 0)

        fields_data = group[['Field Name (MRV)', 'Acres', 'Commodity']].copy()
        fields_data.rename(columns={'Field Name (MRV)': 'Field'}, inplace=True)

        reductions_data = group[['Field Name (MRV)', 'reduced' if batch == "Quantified" else 'reduced_adjusted']].copy()
        reductions_data.rename(columns={'Field Name (MRV)': 'Field'}, inplace=True)
        total_reductions = reductions_data[reductions_data.columns[1]].sum()

        removals_data = group[['Field Name (MRV)', 'baseline_dsoc', 'dsoc', 'removed']].copy()
        removals_data.rename(columns={'Field Name (MRV)': 'Field', 'removed': 'Removed'}, inplace=True)
        total_removals = removals_data['Removed'].sum()
        total_carbon_g = total_reductions + total_removals

        add_custom_bullet(document, "Total Carbon Impacts Generated: ", f"{total_carbon_g:.3f} metric tonnes of carbon dioxide equivalent (mtCO2e)")
        add_custom_bullet(document, "Emission Reductions: ", f"Your practice changes avoided {total_reductions:.3f} mtCO2e.")
        add_custom_bullet(document, "Carbon Removals: ", f"Your soils removed {total_removals:.3f} mtCO2e")
        add_custom_bullet(document, "Total Modeled Acres*: ", f"{fields_data['Acres'].sum():.3f}")

        summary_para5 = document.add_paragraph()
        run_regular = summary_para5.add_run("*ESMC requires complete historical data to generate a baseline, so fields with historical data gaps are not modeled.")
        run_regular.font.name = 'Calibri'
        run_regular.font.size = Pt(9)
        run_regular.font.color.rgb = RGBColor(0, 0, 0)

        para = document.add_heading()
        run = para.add_run(f"Eco-Harvest: rewarding producers for regenerative ag ")
        run_font = run.font
        run_font.name = 'Calibri'
        run_font.size = Pt(16)
        run_font.underline = WD_UNDERLINE.SINGLE
        run_font.color.rgb = RGBColor(0, 0, 0)

        add_custom_bullet(document, "Customizable practices: ", f"Choose how to implement regenerative practices.")
        add_custom_bullet(document, "Fair Payments: ", f"As a non-profit, ESMC maximizes producer benefits.")
        paragraph = document.add_paragraph()
        paragraph.add_run("Many farmers experience a 15-25% return on investment after 3–5 years of ").bold = False
        add_hyperlink(paragraph, 'regenerative farming', 'https://www.bcg.com/publications/2023/regenerative-agriculture-benefits-germany-beyond')
        paragraph.add_run(".").bold = False
        add_custom_bullet(document, "Weather-Ready Farming: ", f"Thank you for participating in the Eco-Harvest program. Continue working with your advisor to improve your soil health and boost your farm’s productivity and earnings.")

        summary_para6 = document.add_paragraph()
        run_bold = summary_para6.add_run('Need Assistance? ')
        run_bold.font.bold = True
        run_bold.font.name = 'Calibri'
        run_bold.font.size = Pt(12)
        run_bold.font.color.rgb = RGBColor(0, 0, 0)
        run_regular = summary_para6.add_run("Contact your Project Representative or ")
        run_regular.font.name = 'Calibri'
        run_regular.font.size = Pt(12)
        run_regular.font.color.rgb = RGBColor(0, 0, 0)
        add_hyperlink(summary_para6, 'support@ecoharvest.ag', 'http://www.ecoharvest.ag')

        # Payment logic
        specific_projects_16 = [
            "AGI SGP Market", "AGI SGP Pipeline", "NACD SGP Market",
            "Northern Plains Cropping", "NACD NGP Market", "NACD NGP Pipeline"
        ]
        specific_projects_20 = ["TNC Minnesota", "TNC Nebraska"]

        payment_amount = 0
        payment_message = "Producers in this project received a payment of $X/acre OR $X/impact."

        if project in specific_projects_16:
            payment_amount = 16 * total_carbon_g
            if payment_amount < 0:
                payment_message = "These fields did not accrue enough impact for payment this growing season."
            else:
                payment_message = f"Producers in this project received a payment of $16/outcome. You will be paid ${payment_amount:.2f}."
        elif project in specific_projects_20:
            payment_amount = 20 * fields_data['Acres'].sum()
            if payment_amount < 0:
                payment_message = "These fields did not accrue enough impact for payment this growing season."
            else:
                payment_message = f"Producers in this project received a payment of $20/acre. You will be paid ${payment_amount:.2f}."

        document.add_paragraph(payment_message)

        # soil sampling data 

        document.add_heading("Soil Sampling Results", level=2)
        soil_bullet_points = [
            "SOC – The percentage of soil that is organic carbon.  2-6% is optimal.",
            "pH – A measurement of acidity or alkalinity of the soil. Optimal rankings are 4-8.",
            "BD – Bulk Density is the oven dry weight of soil per unit volume (g/cm3)",
            "Clay - The percentage of your soil that is clay, 20-35% is optimal."
        ]
        for point in soil_bullet_points:
            paragraph = document.add_paragraph()
            run_bullet_regular = paragraph.add_run(f'• {point}')
            run_bullet_regular.font.name = font_name
            run_bullet_regular.font.size = font_size
            run_bullet_regular.font.color.rgb = font_color

        soil_data = group[['Field Name (MRV)', 'soil_avg_soc', 'soil_avg_ph', 'soil_avg_bulkdensity', 'soil_clay_fraction']].copy()
        soil_data.rename(columns={'Field Name (MRV)': 'Field', 'soil_avg_soc': 'SOC', 'soil_avg_ph': 'pH', 'soil_avg_bulkdensity': 'BD', 'soil_clay_fraction': 'Clay'}, inplace=True)
        create_table(document, soil_data.values, ['Field', 'SOC', 'pH', 'BD', 'Clay'], font_name, font_size, font_color)

        document.add_heading("Carbon Impact Explained", level=2)
        summary_para7 = document.add_paragraph()
        run_regular = summary_para7.add_run("ESMC uses historical data, a scientific model, soil samples and weather data to calculate emissions reductions and soil carbon removals. Any negative emissions are increased emissions, which is not unusual in the early years of adopting a practice change. If concerned, please speak with your technical advisor about further modifications to your practice change(s).")
        run_regular.font.name = 'Calibri'
        run_regular.font.size = Pt(12)
        run_regular.font.color.rgb = RGBColor(0, 0, 0)

        add_custom_bullet(document, "Greenhouse Gas Emissions (GHG) Reductions (mtCO2e): ", f"are the sum of the reductions of methane (CH4), direct nitrous oxide (N2O), indirect nitrous oxide (N2O) and supply chain and operational emissions** between the baseline year and the project year.")
        add_custom_bullet(document, "Soil Carbon Removals (mtCO2e): ", f"The difference in the amount of carbon in the soil between the baseline year and the project year.")

        summary_para8 = document.add_paragraph()
        run_regular = summary_para8.add_run("**Upstream and on-field process emissions associated with management such as fuel used on farm and imbedded emissions of fertilizer products; ESMC utilized emission factors from Ecoinvent and WFLD.")
        run_regular.font.name = 'Calibri'
        run_regular.font.size = Pt(9)
        run_regular.font.color.rgb = RGBColor(0, 0, 0)

        group['Delta Direct N2O'] = calculate_delta(group, 'baseline_n20_direct', 'n2o_direct')
        n2o_data = group[['Field Name (MRV)', 'baseline_n20_direct', 'n2o_direct', 'Delta Direct N2O']].copy()
        n2o_data.rename(columns={
            'Field Name (MRV)': 'Field',
            'baseline_n20_direct': 'Direct N2O Baseline',
            'n2o_direct': 'Direct N2O Practice',
            'Delta Direct N2O': 'Delta'
        }, inplace=True)
        total_delta_direct_n2o = n2o_data['Delta'].sum()

        group['Delta Indirect N2O'] = calculate_delta(group, 'baseline_n2o_indirect' if batch == "Quantified" else 'baseline_n2o_indirect_adjusted', 'n2o_indirect' if batch == "Quantified" else 'n2o_indirect_adjusted')
        n2oindirect_data = group[['Field Name (MRV)', 'baseline_n2o_indirect' if batch == "Quantified" else 'baseline_n2o_indirect_adjusted', 'n2o_indirect' if batch == "Quantified" else 'n2o_indirect_adjusted', 'Delta Indirect N2O']].copy()
        n2oindirect_data.rename(columns={
            'Field Name (MRV)': 'Field',
            'baseline_n2o_indirect' if batch == "Quantified" else 'baseline_n2o_indirect_adjusted': 'Indirect N2O Baseline',
            'n2o_indirect' if batch == "Quantified" else 'n2o_indirect_adjusted': 'Indirect N2O Practice',
            'Delta Indirect N2O': 'Delta'
        }, inplace=True)
        total_delta_indirect_n2o = n2oindirect_data['Delta'].sum()

        group['Delta Methane'] = calculate_delta(group, 'baseline_methane' if batch == "Quantified" else 'baseline_methane_adjusted', 'methane' if batch == "Quantified" else 'methane_adjusted')
        methane_data = group[['Field Name (MRV)', 'baseline_methane' if batch == "Quantified" else 'baseline_methane_adjusted', 'methane' if batch == "Quantified" else 'methane_adjusted', 'Delta Methane']].copy()
        methane_data.rename(columns={
            'Field Name (MRV)': 'Field',
            'baseline_methane' if batch == "Quantified" else 'baseline_methane_adjusted': 'CH4 Baseline',
            'methane' if batch == "Quantified" else 'methane_adjusted': 'CH4 Practice',
            'Delta Methane': 'Delta'
        }, inplace=True)
        total_delta_methane = methane_data['Delta'].sum()

        group['Delta Field Emissions'] = calculate_delta(group, 'field_baseline_emissions', 'field_practice_emissions')
        emissions_data = group[['Field Name (MRV)', 'field_baseline_emissions', 'field_practice_emissions', 'Delta Field Emissions']].copy()
        emissions_data.rename(columns={
            'Field Name (MRV)': 'Field',
            'field_baseline_emissions': 'Supply Chain & Operational Emissions Baseline',
            'field_practice_emissions': 'Supply Chain & Operational Emissions Practice',
            'Delta Field Emissions': 'Delta'
        }, inplace=True)
        total_delta_emissions = emissions_data['Delta'].sum() * 0.001

        impact_table_data = [
            ("Greenhouse Gas Emissions Reductions", total_reductions),
            ("N2O Indirect", total_delta_indirect_n2o),
            ("N2O Direct", total_delta_direct_n2o),
            ("Methane", total_delta_methane),
            ("Supply Chain and Operational Emissions", total_delta_emissions),
            ("Total Soil Carbon Removals", total_removals),
            ("Total Carbon Impact (Reductions + Removals)", total_carbon_g)
        ]
        create_table(document, impact_table_data, ['Impact Breakdown', 'Change between Baseline and practice years (mtCO2e)'], font_name, font_size, font_color)

        para = document.add_heading()
        run = para.add_run(f"Fields Results")
        run_font = run.font
        run_font.name = 'Calibri'
        run_font.size = Pt(14)
        run_font.underline = WD_UNDERLINE.SINGLE
        run_font.color.rgb = RGBColor(112, 173, 71)

        para = document.add_heading()
        run = para.add_run(f"Reductions")
        run_font = run.font
        run_font.name = 'Calibri'
        run_font.size = Pt(12)

        document.add_paragraph("This section shows the GHG emissions reductions from the intervention(s) per field. Total reductions include the following emissions: direct nitrogen, methane and supply chain & operational emissions.")

        document.add_paragraph()
        total_reductions = reductions_data[reductions_data.columns[1]].sum()
        para = document.add_paragraph()
        run = para.add_run(f"Total Reductions: {total_reductions:.3f} tonnes CO2e")
        run.italic = True

        create_table(document, reductions_data.values, ['Field', 'Total Reductions'], font_name, font_size, font_color)

        document.add_paragraph()
        para = document.add_paragraph()
        run = para.add_run("Direct Nitrogen (N2O) Emissions")
        run.bold = True
        run.font.name = 'Calibri'

        document.add_paragraph()
        total_n2o_reductions = n2o_data['Delta'].sum()
        para = document.add_paragraph()
        run = para.add_run(f"Total Direct N2O Reductions: {total_n2o_reductions:.3f} tonnes CO2e")
        run.italic = True

        create_table(document, n2o_data.values, ['Field', 'Direct N2O Baseline', 'Direct N2O Practice', 'Delta'], font_name, font_size, font_color)

        document.add_paragraph()
        para = document.add_paragraph()
        run = para.add_run("Indirect Nitrogen (N2O) Emissions")
        run.bold = True
        run.font.name = 'Calibri'

        document.add_paragraph()
        total_indirectn2o_reductions = n2oindirect_data['Delta'].sum()
        para = document.add_paragraph()
        run = para.add_run(f"Total Indirect N2O Reductions: {total_indirectn2o_reductions:.3f} tonnes CO2e")
        run.italic = True

        create_table(document, n2oindirect_data.values, ['Field', 'Indirect N2O Baseline', 'Indirect N2O Practice', 'Delta'], font_name, font_size, font_color)

        document.add_paragraph()
        para = document.add_paragraph()
        run = para.add_run("Methane (CH4) Emissions")
        run.bold = True
        run.font.name = 'Calibri'

        document.add_paragraph()
        total_ch4_reductions = methane_data['Delta'].sum()
        para = document.add_paragraph()
        run = para.add_run(f"Total CH4 Reductions: {total_ch4_reductions:.3f} tonnes CO2e")
        run.italic = True

        create_table(document, methane_data.values, ['Field', 'CH4 Baseline', 'CH4 Practice', 'Delta'], font_name, font_size, font_color)

        document.add_paragraph()
        para = document.add_paragraph()
        run = para.add_run("Supply Chain and Operational Emissions")
        run.bold = True
        run.font.name = 'Calibri'

        document.add_paragraph()
        total_emissions_reductions = emissions_data['Delta'].sum()
        para = document.add_paragraph()
        run = para.add_run(f"Total Supply Chain and Operational Emissions: {total_delta_emissions:.3f} tonnes CO2e")
        run.italic = True

        create_table(document, emissions_data.values, ['Field', 'Supply Chain & Operational Emissions Baseline', 'Supply Chain & Operational Emissions Practice', 'Delta'], font_name, font_size, font_color, scale_emissions=True)

        para = document.add_heading()
        run = para.add_run(f"Removals")
        run_font = run.font
        run_font.name = 'Calibri'
        run_font.size = Pt(12)

        document.add_paragraph("This section shows the modeled carbon removals per field from the intervention(s).")

        document.add_paragraph()
        para = document.add_paragraph()
        run = para.add_run(f"Total Reductions: {total_removals:.3f} tonnes CO2e")
        run.italic = True

        removals_data.rename(columns={'baseline_dsoc': 'Removals Baseline', 'dsoc': 'Removals Practice', 'Removed': 'Delta'}, inplace=True)
        create_table(document, removals_data.values, ['Field', 'Removals Baseline', 'Removals Practice', 'Delta'], font_name, font_size, font_color)

        para = document.add_heading()
        run = para.add_run(f"Payment Structure")
        run_font = run.font
        run_font.name = 'Calibri'
        run_font.size = Pt(12)

        document.add_paragraph(payment_message)

        document.add_picture('diagram.png')
        italic_paragraph = document.add_paragraph()
        italic_run = italic_paragraph.add_run('Figure 1: Total impact is the combined value of GHG emissions reductions and soil carbon stored. This diagram illustrates a general improvement from the baseline to project year – GHG emissions decrease and soil carbon increases.')
        italic_run.italic = True
        italic_run.font_name = 'Calibri'
        italic_run.font.size = Pt(11)

        document.add_paragraph()
        para = document.add_paragraph()
        run = para.add_run("Real-world Comparisons")
        run.bold = True
        run.font.name = 'Calibri'

        summary_para9 = document.add_paragraph()
        run_regular = summary_para9.add_run("Use ")
        run_regular.font.name = 'Calibri'
        run_regular.font.size = Pt(12)
        run_regular.font.color.rgb = RGBColor(0, 0, 0)
        add_hyperlink(summary_para9, 'this tool', 'https://www.epa.gov/energy/greenhouse-gas-equivalencies-calculator#resultsg')
        run_regular = summary_para9.add_run(" to compare carbon impacts to everyday examples.")
        run_regular.font.name = 'Calibri'
        run_regular.font.size = Pt(12)
        run_regular.font.color.rgb = RGBColor(0, 0, 0)

    return doc.producers

# Load data
verified_data = load_data(run_year, "Verified")
quantified_data = load_data(run_year, "Quantified")

# Initialize document dictionary
class DocumentHolder:
    def __init__(self):
        self.producers = {}

doc = DocumentHolder()

# Generate reports for each batch
generate_report(verified_data, "Verified", doc)
generate_report(quantified_data, "Quantified", doc)

# Save documents
output_dir = "Producers"
if not os.path.exists(output_dir):
    os.makedirs(output_dir)

for producer, document in doc.producers.items():
    display_name = producer_mapping.get(producer, producer)  # Ensure you get the display name for each producer
    try:
        safe_producer_name = producer.replace('/', '_').replace('\\', '_')
        safe_display_name = display_name.replace('/', '_').replace('\\', '_')
        doc_file_name = f"{safe_display_name} ({safe_producer_name}).docx"
    except:
        safe_producer_name = producer.replace('/', '_').replace('\\', '_')
        doc_file_name = f"{safe_producer_name}.docx"
    document.save(os.path.join(output_dir, doc_file_name))


print("Reports generated successfully.")

import pandas as pd
from docx import Document
from docx.shared import Inches, Pt
from docx.oxml import OxmlElement
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.text import WD_UNDERLINE
from docx.shared import RGBColor
from docx.oxml.ns import qn
import os

# Load the data
data_file = 'Producer Field Soil Data.csv'
data = pd.read_csv(data_file)

# Convert Soil Sample Year to integer
data['Soil Sample Year'] = data['Soil Sample Year'].astype('Int64')

# List of projects to ignore
ignore_projects = [
    "All Scopes and Assets_Credits", "Austin-Test-1", "Austin-Test-3", 
    "Cabbage Patch- Will Test", "Data Modeling", "ESMC-Test", 
    "G's EoY test project", "Laura Test", "Michelle-test", 
    "Producer Circle", "Regrow test project", "SLM Demo", 
    "The Fertilizer Institute (TFI)", "Travis-Test1", "Test Project 10", "Historical Fields",
    "CIF Test Project", "ESMC-PM-Test"
]

# Filter out the ignored projects
data = data[~data['Project'].isin(ignore_projects)]

def format_cell_data(data):
    try:
        float_data = float(data)
        return f"{float_data:.3f}"
    except ValueError:
        return data

def create_table(doc, data, headers, font_name, font_size, font_color):
    table = doc.add_table(rows=len(data) + 1, cols=len(headers))
    table.style = 'Table Grid'

    for i, header in enumerate(headers):
        cell = table.cell(0, i)
        run = cell.paragraphs[0].add_run(header)
        run.font.name = font_name
        run.font.size = font_size
        run.font.color.rgb = font_color
        run.font.bold = True
        cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    for row_idx, row_data in enumerate(data, start=1):
        for col_idx, cell_data in enumerate(row_data):
            cell_data = format_cell_data(cell_data)
            cell = table.cell(row_idx, col_idx)
            run = cell.paragraphs[0].add_run(str(cell_data))
            run.font.name = font_name
            run.font.size = font_size
            run.font.color.rgb = font_color
            cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

def add_custom_bullet(doc, label, value):
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(label)
    run.bold = True
    run.font.name = 'Calibri'
    run.font.size = Pt(12)

    value_run = p.add_run(value)
    value_run.font.name = 'Calibri'
    value_run.font.size = Pt(12)

def add_hyperlink(paragraph, text, url):
    part = paragraph.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)

    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)
    hyperlink.set(qn('w:history'), '1')

    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    c = OxmlElement('w:color')
    c.set(qn('w:val'), '0000FF')
    rPr.append(c)
    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'single')
    rPr.append(u)
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)

    return hyperlink

def initialize_document():
    doc = Document()
    # Add header image only once
    section = doc.sections[0]
    header = section.header
    header_para = header.paragraphs[0]
    run = header_para.add_run()
    run.add_picture('header.png', width=Inches(6))
    return doc

def generate_soil_summary_report(data, doc):
    font_name = 'Calibri'
    font_size = Pt(12)
    font_color = RGBColor(0, 0, 0)

    for producer, group in data.groupby('Producer (MRV Name)'):
        if producer not in doc.producers:
            doc.producers[producer] = initialize_document()

        document = doc.producers[producer]

        anonymized_name = group['Producer (Anonymized)'].iloc[0]
        project_name = group['Project'].iloc[0]
        sample_years = group['Soil Sample Year'].dropna().unique()

        document.add_paragraph()  # empty space
        para = document.add_heading()
        run = para.add_run(f"Soil Summary Report")
        run_font = run.font
        run_font.name = 'Calibri'
        run_font.size = Pt(24)
        run_font.underline = WD_UNDERLINE.SINGLE
        run_font.color.rgb = RGBColor(0, 0, 0)

        summary_para1 = document.add_paragraph()
        run_bold = summary_para1.add_run('Producer (MRV Name): ')
        run_bold.font.bold = True
        run_bold.font.name = 'Calibri'
        run_bold.font.size = Pt(12)
        run_bold.font.color.rgb = RGBColor(0, 0, 0)
        run_regular = summary_para1.add_run(producer)
        run_regular.font.name = 'Calibri'
        run_regular.font.size = Pt(12)
        run_regular.font.color.rgb = RGBColor(0, 0, 0)

        summary_para2 = document.add_paragraph()
        run_bold = summary_para2.add_run('Producer (Anonymized): ')
        run_bold.font.bold = True
        run_bold.font.name = 'Calibri'
        run_bold.font.size = Pt(12)
        run_bold.font.color.rgb = RGBColor(0, 0, 0)
        run_regular = summary_para2.add_run(anonymized_name)
        run_regular.font.name = 'Calibri'
        run_regular.font.size = Pt(12)
        run_regular.font.color.rgb = RGBColor(0, 0, 0)

        summary_para3 = document.add_paragraph()
        run_bold = summary_para3.add_run('Project: ')
        run_bold.font.bold = True
        run_bold.font.name = 'Calibri'
        run_bold.font.size = Pt(12)
        run_bold.font.color.rgb = RGBColor(0, 0, 0)
        run_regular = summary_para3.add_run(project_name)
        run_regular.font.name = 'Calibri'
        run_regular.font.size = Pt(12)
        run_regular.font.color.rgb = RGBColor(0, 0, 0)

        summary_para4 = document.add_paragraph()
        run_regular = summary_para4.add_run('ESMC requires complete historical data to model a field. If historical data is incomplete, we are unable to model the field. Additional reasons why were not able to model a field include grazing events, the use of custom fertilizers or compost, crops not supported by the model or tile drainage events. We appreciate your patience as we work with our modeling partners to model fields with these events, crops and nuances. Please check with your implementation partner for further assistance.')
        run_regular.font.name = 'Calibri'
        run_regular.font.size = Pt(12)
        run_regular.font.color.rgb = RGBColor(0, 0, 0)


        for year in sample_years:
            year_group = group[group['Soil Sample Year'] == year]
            if len(sample_years) > 1:
                document.add_heading(f"Soil Sample Year: {int(year)}", level=2)

            fields_data = year_group[['Field (MRV Name)', 'acres', 'sampled points', 'soc_avg', 'bd_avg']].copy()
            fields_data.rename(columns={
                'Field (MRV Name)': 'Field',
                'acres': 'Acres',
                'sampled points': 'Sampled Points',
                'soc_avg': 'SOC Avg',
                'bd_avg': 'BD Avg'
            }, inplace=True)

            create_table(document, fields_data.values, ['Field', 'Acres', 'Sampled Points', 'SOC Avg', 'BD Avg'], font_name, font_size, font_color)

    return doc.producers

# Initialize document dictionary
class DocumentHolder:
    def __init__(self):
        self.producers = {}

doc = DocumentHolder()

# Generate reports
generate_soil_summary_report(data, doc)

# Save documents
output_dir = "Soil_Summary_Reports"
if not os.path.exists(output_dir):
    os.makedirs(output_dir)

for producer, document in doc.producers.items():
    group = data[data['Producer (MRV Name)'] == producer]
    project_name = group['Project'].iloc[0]
    safe_project_name = project_name.replace('/', '_').replace('\\', '_')

    project_dir = os.path.join(output_dir, safe_project_name)
    if not os.path.exists(project_dir):
        os.makedirs(project_dir)

    safe_producer_name = producer.replace('/', '_').replace('\\', '_')
    doc_file_name = f"{safe_producer_name}.docx"
    document.save(os.path.join(project_dir, doc_file_name))

print("Soil summary reports generated successfully.")
